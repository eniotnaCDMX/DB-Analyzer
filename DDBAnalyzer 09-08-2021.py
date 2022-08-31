# -*- coding: utf-8 -*-
from math import *
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import tkinter as tk
from tkinter import ttk
from tkinter.filedialog import askopenfilename
import tkinter.font as font
import os
import os.path
import csv
import docx
from docx.shared import Inches
import matplotlib.pyplot as plt
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx.enum.text import WD_ALIGN_PARAGRAPH

##### Fichero actualizado el 09/08/2021 #####

# Este programma nomas permite tratar datos del tipo Excel datos
TN = 0
Icc = 0
Inom = 0
Snom = 0
FPQ = 0
# Permite conseguir el camino del desktop de la computadora
def get_desktop():    
    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop') 
    return desktop

# Nueva carpeta creada en el desktop con el nombre : Nuevo Estudio
# Una vez creada, por favor, cambiar el nombre de la carpeta
if os.path.exists(get_desktop() + "\\Nuevo Estudio") == False:     
    dossier = os.makedirs(get_desktop() + "\\Nuevo Estudio") 

# Este crea un documiento en .docx que va a contener nuestro estudio
document = docx.Document(get_desktop() + "\\DDBAnalyzer" + "\\Template PQB.docx")

# Empezamos a editar un poco el documiento .docx
document.add_heading('Estudio PQ Barcon numero:', 0)
document.add_heading("Introducción:",1)
document.add_paragraph("PQ Barcon es una empresa mexicana creada por un grupo de empresarios en el sector eléctrico mexicano apasionados por la ingeniería, que deciden apoyar al Código de Red que se implanta en México para Centros de Carga a partir de abril del 2019." + "\n")
document.add_paragraph("Este documiento tiene por objectivo de analyzar una serie de medidas cinco minutales "\
                       + "realisadas 24/24 durante una semana, y de presentar estos datos de forma gráfica"\
                       + " acompañada con indicadores estadisticos y con comparación a las normas y recomendaciónes" \
                       + " que se aplican en este campo." + "\n")
document.add_paragraph("Especificamente, este documiento contendra los siguientes datos:" + "\n")
document.add_paragraph("Parte 1: Potencia activa, aparente, reactiva y de distorsión.")
document.add_paragraph("Parte 2: Frecuencia.")
document.add_paragraph("Parte 3: Tensiones y corrientes.")
document.add_paragraph("Parte 4: Desbalance de tensión y de corriente.")
document.add_paragraph("Parte 5: Flicker y distorción de señales.")
document.add_paragraph("Parte 6: Factor de potencia" + "\n")
document.add_paragraph("En anexo, se ve explicado los metodos de calculo y sus justificaciones. ")


# Permite conseguir los datos de un fichero .csv y ponerlos en una lista
def extractiondonne(fichier):   # Révisé !!! # Normalement version définitive
    Liste = []
    with open(fichier,'r') as fst:
        lecteur = csv.reader(fst,delimiter = ';')
        for ligne in lecteur:
            Liste.append(ligne)
        for ligne in range(1,len(Liste)):
            for colonne in range(1,len(Liste[0])):
                Liste[ligne][colonne] = float(Liste[ligne][colonne].replace(",","."))        
    return Liste

    
# En una ventana tkinter, voltea al valor promedio de Pavg, Pmin y Pmax, a sus valores medianas, al los valores
# maxima y minima encontradas y en otra ventana, desplega al grafico de Pavg, Pmin y Pmax, y a la media mobil 
# sobre 6 horas         
def BoutonP(fichier):  # Révisé !!!
    L = extractiondonne(fichier) 
    Pmoy = []
    Pmax = []
    Pmin = []
    NB = []
    for i in range(1,len(L)): # Contient les indices des différentes lignes    
        NB.append(i)
    for i in range(1,len(L)): # LLena una lista con los valores de la potencia activa promedia
        Pmoy.append(L[i][47])    
    for i in range(1,len(L)): # Llena una lista con los valores de la potencia activa maxima
        Pmax.append(L[i][46])
    for i in range(1,len(L)): # Llena una lista con los valores de la potencia activa minima
        Pmin.append(L[i][48])
    FenetreP = tk.Tk()
    FenetreP.wm_title("Potencia activa (kW)")
    FenetreP.configure(background="#2B00FA")
    fig = Figure(figsize=(6, 4), dpi=96)
    ax = fig.add_subplot(111)
    ax.plot(NB,Pmoy, label = "Ppro")
    ax.legend()
    ax.plot(NB,Pmax, label = "Pmax")
    ax.legend()
    ax.plot(NB,Pmin, label = "Pmin")
    ax.legend()
    ax.plot(NB,moyenneGli(Pmoy,36), label = "avg 6horas")
    ax.legend()
    graph = FigureCanvasTkAgg(fig, master=FenetreP)
    canvas = graph.get_tk_widget()
    canvas.grid(row=0, column=0)
    TexteP = tk.Tk()
    TexteP.wm_title("Potencia activa (kW)")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1 = tk.Label(TexteP, text = "<Pprom> = " + str(ceil(moyenneA(Pmoy))) + " kW" + "  "\
                      + "<Pmax> =  "\
                     + str(ceil(moyenneA(Pmax))) + " kW" + "  " +\
                         "<Pmin> = "\
                     + str(ceil(moyenneA(Pmin))) + " kW", foreground = "white", background = "#2B00FA")
    label1['font'] = f_label
    label1.pack()
    label2 = tk.Label(TexteP, text= "Pprom mediana = " + str(ceil(mediane(Pmoy))) + " kW" + "  " +\
                                    "Pmax mediana = " + str(ceil(mediane(Pmax))) + " kW" + "  " +\
                                    "Pmin mediana = " + str(ceil(mediane(Pmin))) + " kW",foreground = "white", background = "#2B00FA")
                                                                           
    label2['font'] = f_label
    label2.pack()
    label3 = tk.Label(TexteP, text="La potencia maxima encontrada vale "\
                     + str(ceil(max(Pmax))) + " kW" + "\n" + "La potencia minima encontrada vale "\
                     + str(ceil(min(Pmin))) + " kW", foreground = "white", background = "#2B00FA")
    label3['font'] = f_label
    label3.pack()

# Esta funcion permite conseguir los datos de la potencia activa maxima, minima y promedia
# y escribe adentro del Word la parte sobre la potencia activa, osea :
# escribe el valor promedio, la mediana, la potencia maxima y minima encontrada
# y desplegue un grafico donde hay Ppro, Pmax y Pmin     
def P(fichier): # Révisé !!!
    L = extractiondonne(fichier)    
    Pmoy = []
    Pmax = []
    Pmin = []
    NB = []
    for i in range(1,len(L)): # Contient les indices des différentes lignes    
        NB.append(i)    
    for i in range(1,len(L)):
        Pmoy.append(L[i][47])    
    for i in range(1,len(L)):
        Pmax.append(L[i][46])
    for i in range(1,len(L)):
        Pmin.append(L[i][48])
    plt.plot(NB,Pmoy, label = r"$P_{pro}$")
    plt.legend()
    plt.plot(NB,Pmax, label = r"$P_{max}$")
    plt.legend()
    plt.plot(NB,Pmin, label = r"$P_{min}$")
    plt.legend()
    plt.plot(NB,moyenneGli(Pmoy,36), label = "$P_{{pro}_{6horas}}$")
    plt.legend()        
    plt.title("Potencia activa (kW)")
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Potencia activa.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Potencia activa.png")
        plt.clf()
    document.add_heading("Partida 1: potencias en el sistema ",1)
    document.add_heading("1.1 Potencia activa: ",2)
    document.add_paragraph("La potencia activa en una fase se define cómo el valor promedio en un periodo del"\
                           + " producto entre el corriente en esa fase y la tensión entre esa fase y el neutro"\
                           + " osea:" + "\n")
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\definition P.png", width=Inches(2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER                       
    document.add_paragraph("Este grafico mostra la potencia activa en el sistema por cada medida: " + "\n")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Potencia activa.png", width=Inches(5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio de la potencia activa promedia vale {:.0f}".format(moyenneA(Pmoy)) + " kW.")
    document.add_paragraph("- El valor promedio de la potencia activa maxima vale {:.0f}".format(moyenneA(Pmax)) + " kW.")                       
    document.add_paragraph("- El valor promedio de la potencia activa minima vale {:.0f}".format(moyenneA(Pmin)) + " kW.")
    document.add_paragraph("- La mediana de la potencia activa promedia vale {:.0f}".format(mediane(Pmoy)) + " kW.")
    document.add_paragraph("- La mediana de la potencia activa maxima vale {:.0f}".format(mediane(Pmax)) + " kW.")                       
    document.add_paragraph("- La mediana de la potencia activa minima vale {:.0f}".format(mediane(Pmin)) + " kW.")
    document.add_paragraph("- La potencia maxima encontrada es de {:.0f}".format(max(Pmax)) + " kW.")
    document.add_paragraph("- La potencia minima encontrada es de {:.0f}".format(min(Pmin)) + " kW." + "\n")

# En una ventana tkinter, voltea al valor promedio de Spro, Smin y Smax, a sus valores medianas, al los valores
# maxima y minima encontradas y en otra ventana, desplega al grafico de Ppro, Pmin y Pmax, y a la media mobil 
# sobre 6 horas         
def BoutonS(fichier):  # Révisé !!! #A améliorer éventuellement
    L = extractiondonne(fichier) 
    Pmoy = []
    Pmax = []
    Pmin = []
    NB = []
    for i in range(1,len(L)): # Contient les indices des différentes lignes    
        NB.append(i)
    for i in range(1,len(L)): # LLena una lista con los valores de la potencia activa promedia
        Pmoy.append(L[i][95])    
    for i in range(1,len(L)): # Llena una lista con los valores de la potencia activa maxima
        Pmax.append(L[i][94])
    for i in range(1,len(L)): # Llena una lista con los valores de la potencia activa minima
        Pmin.append(L[i][96])
    FenetreP = tk.Tk()
    FenetreP.wm_title("Potencia aparente (kVA)")
    FenetreP.configure(background="#2B00FA")
    fig = Figure(figsize=(6, 4), dpi=96)
    ax = fig.add_subplot(111)
    ax.plot(NB,Pmoy, label = "Spro")
    ax.legend()
    ax.plot(NB,Pmax, label = "Smax")
    ax.legend()
    ax.plot(NB,Pmin, label = "Smin")
    ax.legend()
    ax.plot(NB,moyenneGli(Pmoy,36), label = "avg 6horas")
    ax.legend()
    graph = FigureCanvasTkAgg(fig, master=FenetreP)
    canvas = graph.get_tk_widget()
    canvas.grid(row=0, column=0)
    TexteP = tk.Tk()
    TexteP.wm_title("Potencia activa (kW)")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1 = tk.Label(TexteP, text = "<Sprom> = " + str(ceil(moyenneA(Pmoy))) + " kVA" + "  "\
                      + "<Smax> =  "\
                     + str(ceil(moyenneA(Pmax))) + " kVA" + "  " +\
                         "<Smin> = "\
                     + str(ceil(moyenneA(Pmin))) + " kVA", foreground = "white", background = "#2B00FA")
    label1['font'] = f_label
    label1.pack()
    label2 = tk.Label(TexteP, text= "Sprom mediana = " + str(ceil(mediane(Pmoy))) + " kVA" + "  " +\
                                    "Smax mediana = " + str(ceil(mediane(Pmax))) + " kVA" + "  " +\
                                    "Smin mediana = " + str(ceil(mediane(Pmin))) + " kVA",foreground = "white", background = "#2B00FA")
                                                                           
    label2['font'] = f_label
    label2.pack()
    label3 = tk.Label(TexteP, text="La potencia maxima encontrada vale "\
                     + str(ceil(max(Pmax))) + " kVA" + "\n" + "La potencia minima encontrada vale "\
                     + str(ceil(min(Pmin))) + " kVA", foreground = "white", background = "#2B00FA")
    label3['font'] = f_label
    label3.pack()

def S(fichier): # Révisé !!! #A modifier éventuellement
    L = extractiondonne(fichier)    
    Pmoy = []
    Pmax = []
    Pmin = []
    NB = []
    for i in range(1,len(L)): # Contient les indices des différentes lignes    
        NB.append(i)    
    for i in range(1,len(L)):
        Pmoy.append(L[i][95])    
    for i in range(1,len(L)):
        Pmax.append(L[i][94])
    for i in range(1,len(L)):
        Pmin.append(L[i][96])
    plt.plot(NB,Pmoy, label = r"$S_{pro}$")
    plt.legend()
    plt.plot(NB,Pmax, label = r"$S_{max}$")
    plt.legend()
    plt.plot(NB,Pmin, label = r"$S_{min}$")
    plt.legend()
    plt.plot(NB,moyenneGli(Pmoy,36), label = r"$S_{{pro}_{6horas}}$")
    plt.legend()        
    plt.title("Potencia aparente (kVA)")
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Potencia aparente.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Potencia aparente.png")
        plt.clf()
    document.add_heading("1.2 Potencia aparente: ",2)
    document.add_paragraph("La potencia aparente en una fase se define cómo el producto entre el valor eficaz"\
                           + " de la tension de fase y el valor eficaz del coriente en esa fase, o sea:" + "\n")    
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\definition S.png", width=Inches(1.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("\n" + "Este grafico mostra la evolucion de la potencia aparente: ")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Potencia aparente.png", width=Inches(5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio de la potencia aparente promedia vale {:.0f}".format(moyenneA(Pmoy)) + " kVA.")
    document.add_paragraph("- El valor promedio de la potencia aparente maxima vale {:.0f}".format(moyenneA(Pmax)) + " kVA.")                       
    document.add_paragraph("- El valor promedio de la potencia aparente minima vale {:.0f}".format(moyenneA(Pmin)) + " kVA.")
    document.add_paragraph("- La mediana de la potencia aparente promedia vale {:.0f}".format(mediane(Pmoy)) + " kVA.")
    document.add_paragraph("- La mediana de la potencia aparente maxima vale {:.0f}".format(mediane(Pmax)) + " kVA.")                       
    document.add_paragraph("- La mediana de la potencia aparente minima vale {:.0f}".format(mediane(Pmin)) + " kVA.")
    document.add_paragraph("- La potencia maxima encontrada es de {:.0f}".format(max(Pmax)) + " kVA.")
    document.add_paragraph("- La potencia minima encontrada es de {:.0f}".format(min(Pmin)) + " kVA." + "\n")

def BoutonQ(fichier):
    L = extractiondonne(fichier) 
    Pmoy = []
    NB = []
    for i in range(1,len(L)): # Contient les indices des différentes lignes    
        NB.append(i)
    for i in range(1,len(L)): # LLena una lista con los valores de la potencia reactiva promedia
        Pmoy.append(L[i][103])    
    plt.plot(NB,Pmoy, label = r"$Q_{pro}$")
    plt.legend()
    plt.plot(NB,moyenneGli(Pmoy,36), label = r"$Q_{{pro}_{6horas}}$")
    plt.legend()        
    plt.title("Potencia reactiva antecompensacion (kVAR)")    
    FenetreP = tk.Tk()
    FenetreP.wm_title("Potencia reactiva ante compensacion (kVAR)")
    FenetreP.configure(background="#2B00FA")
    fig = Figure(figsize=(6, 4), dpi=96)
    ax = fig.add_subplot(111)
    ax.plot(NB,Pmoy, label = "Qmoy")
    ax.legend()
    ax.plot(NB,moyenneGli(Pmoy,36), label = "prom 6horas")
    ax.legend()
    graph = FigureCanvasTkAgg(fig, master=FenetreP)
    canvas = graph.get_tk_widget()
    canvas.grid(row=0, column=0)
    TexteP = tk.Tk()
    TexteP.wm_title("Potencia reactiva ante compensacion")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1 = tk.Label(TexteP, text = "<Qprom> = {:.0f}".format(moyenneA(Pmoy)) + " kVAR",foreground = "white", background = "#2B00FA") 
    label1['font'] = f_label
    label1.pack()
    label2 = tk.Label(TexteP, text= "Qprom mediana = {:.0f}".format(mediane(Pmoy)) + " kVAR",foreground = "white", background = "#2B00FA")
                                                                           
    label2['font'] = f_label
    label2.pack()
    label3 = tk.Label(TexteP, text="La potencia maxima encontrada vale {:.0f}".format(max(Pmoy))\
                     + " kVAR" + "\n" + "La potencia minima encontrada vale {:.0f}".format(min(Pmoy)) + " kVAR", foreground = "white", background = "#2B00FA")
    label3['font'] = f_label
    label3.pack()

def Q(fichier): # P = 47, fp = 89, ya !
    L = extractiondonne(fichier) 
    Pmoy = []
    NB = []
    Qc = []
    for i in range(1,len(L)): # Contient les indices des différentes lignes    
        NB.append(i)
    for i in range(1,len(L)): # LLena una lista con los valores de la potencia reactiva promedia
        Pmoy.append(L[i][103])
        Qc.append(L[i][47]*(tan(acos(L[i][89])) - tan(acos(FPQ))))
    plt.plot(NB,Pmoy, label = r"$Q_{pro}$")
    plt.legend()
    plt.plot(NB,moyenneGli(Pmoy,36), label = r"$Q_{{pro}_{6horas}}$")
    plt.legend()        
    plt.title("Potencia reactiva antecompensacion (kVAR)")    
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Potencia reactiva AC.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Potencia reactiva AC.png")
        plt.clf()
    plt.plot(NB,Qc, label = r"$Qc$")
    plt.legend()
    plt.plot(NB,moyenneGli(Qc,36), label = r"$Qc_{{pro}_{6horas}}$")
    plt.legend()        
    plt.title("Potencia reactiva de compensación (kVAR)")    
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Potencia reactiva C.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Potencia reactiva C.png")
        plt.clf()
    document.add_heading("1.3 Potencia reactiva: ",2)
    document.add_paragraph("Este grafico mostra la evolución de la potencia reactiva antecompensación: ")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Potencia reactiva AC.png", width=Inches(5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio de la potencia reactiva vale {:.0f}".format(moyenneA(Pmoy)) + " kVAR.")
    document.add_paragraph("- La mediana de la potencia reactiva promedia vale {:.0f}".format(mediane(Pmoy)) + " kVAR.")
    document.add_paragraph("- La potencia maxima encontrada vale {:.0f}".format(max(Pmoy)) + " kVAR.")
    document.add_paragraph("- La potencia minima encontrada vale {:.0f}".format(min(Pmoy)) + " kVAR." + "\n")
    document.add_paragraph("Si quisieramos compensar esa potencia reactiva, con un factor de potencia final de " + str(FPQ)\
                           + " nuestros dispositivos de compensation deberían producir"\
                           + " la potencia reactiva siguiente:"    )
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Potencia reactiva C.png", width=Inches(5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Relevamos los datos siguientes: " + "\n")
    document.add_paragraph("- El valor promedio de la potencia reactiva vale {:.0f}".format(moyenneA(Qc)) + " kVAR.")
    document.add_paragraph("- La mediana de la potencia reactiva promedia vale {:.0f}".format(mediane(Qc)) + " kVAR.")
    document.add_paragraph("- La potencia maxima encontrada vale {:.0f}".format(max(Qc)) + " kVAR.")
    document.add_paragraph("- La potencia minima encontrada vale {:.0f}".format(min(Qc)) + " kVAR." + "\n")
    document.add_paragraph("Con una frecuancia nominal de 60 Hz y una tension nominal de fase de " + str(TN) + " V"\
                           + " si quisieramos compensar esa potencia con bancos de capacitores cablados en triangulo"\
                           + " esos capacitores deberían tener una capacidad: C = {:.1E}".format((mediane(Qc)*1000)/(360*pi*(sqrt(3)*TN)**2)) + " F")
def D(fichier):
    L = extractiondonne(fichier)
    NB = []
    HA = []
    HB = []
    HC = []
    A = []
    b = 0
    for i in range(1,len(L)):
        NB.append(i)
    for i in range(1,len(L)):
        for j in range (261,261 + 47*3 + 1,3): #fase A 255 fase B 256 fase C 257      
            A.append(L[i][j])
        for k in range(len(A)):
            b += A[k]**2
        HA.append(L[i][2]*sqrt(b)/1000)
        A = []
        b = 0
    for i in range(1,len(L)):
        for j in range (262,262 + 3*47 + 1,3): #fase A 255 fase B 256 fase C 257      
            A.append(L[i][j])
        for k in range(len(A)):
            b += A[k]**2
        HB.append(L[i][5]*sqrt(b)/1000)
        A = []
        b = 0
    for i in range(1,len(L)):
        for j in range (263,263 + 3*47 + 1,3): #fase A 255 fase B 256 fase C 257      
            A.append(L[i][j])
        for k in range(len(A)):
            b += A[k]**2
        HC.append(L[i][8]*sqrt(b)/1000)
        A = []
        b = 0    
    plt.plot(NB,HA, label = r"$D_{A}$")
    plt.legend()
    plt.plot(NB,HB, label = r"$D_{B}$")
    plt.legend()
    plt.plot(NB,HC, label = r"$D_{C}$")
    plt.legend()        
    plt.title("Potencia de deformación (kVAD)")
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Potencia deformante.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Potencia deformante.png")
        plt.clf()
    document.add_heading("1.4 Potencia de deformación: ",2)
    document.add_paragraph("La potencia de deformación en una fase se define via las tres potencias encontradas de la manera siguiente: " + "\n")    
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\definition D.png", width=Inches(2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Este grafico mostra la evolucion de la potencia de deformación: " + "\n")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Potencia deformante.png", width=Inches(5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio de D en la fase A vale {:.0f}".format(moyenneA(HA)) + " kVAD.")
    document.add_paragraph("- El valor promedio de D en la fase B vale {:.0f}".format(moyenneA(HB)) + " kVAD.")                       
    document.add_paragraph("- El valor promedio de D en la fase C vale {:.0f}".format(moyenneA(HC)) + " kVAD.")
    document.add_paragraph("- La mediana de D en la fase A vale {:.0f}".format(mediane(HA)) + " kVAD.")
    document.add_paragraph("- La mediana de D en la fase B vale {:.0f}".format(mediane(HB)) + " kVAD.")                       
    document.add_paragraph("- La mediana de D en la fase C vale {:.0f}".format(mediane(HC)) + " kVAD.")
    document.add_paragraph("- La potencia maxima encontrada en la fase A vale {:.0f}".format(max(HA)) + " kVAD.")
    document.add_paragraph("- La potencia maxima encontrada en la fase B vale {:.0f}".format(max(HB)) + " kVAD.")
    document.add_paragraph("- La potencia maxima encontrada en la fase C vale {:.0f}".format(max(HC)) + " kVAD.")
    document.add_paragraph("- La potencia minima encontrada en la fase A vale {:.0f}".format(min(HA)) + " kVAD.")
    document.add_paragraph("- La potencia minima encontrada en la fase B vale {:.0f}".format(min(HB)) + " kVAD.")
    document.add_paragraph("- La potencia minima encontrada en la fase C vale {:.0f}".format(min(HC)) + " kVAD." + "\n")

# Da el corriente maximo de todas las medidas
def GetIM(fichier):
    L = extractiondonne(fichier) 
    IAmax = [] #23
    IBmax = [] #26
    ICmax = [] #29
    for i in range(1,len(L)):
        IAmax.append(L[i][23])
        IBmax.append(L[i][26])
        ICmax.append(L[i][29])
    return max(max(IAmax),max(IBmax),max(ICmax))
    

def BoutonFp(fichier): #Révisé, actualisé via .format
    L = extractiondonne(fichier) 
    Pmoy = []
    Pmax = []
    Pmin = []
    NB = []
    for i in range(1,len(L)): # Contient les indices des différentes lignes    
        NB.append(i)
    for i in range(1,len(L)): # LLena una lista con los valores de la potencia activa promedia
        Pmoy.append(L[i][89])    
    for i in range(1,len(L)): # Llena una lista con los valores de la potencia activa maxima
        Pmax.append(L[i][88])
    for i in range(1,len(L)): # Llena una lista con los valores de la potencia activa minima
        Pmin.append(L[i][90])
    FenetreP = tk.Tk()
    FenetreP.wm_title("Factor de potencia")
    FenetreP.configure(background="#2B00FA")
    fig = Figure(figsize=(6, 4), dpi=96)
    ax = fig.add_subplot(111)
    ax.plot(NB,Pmoy, label = "FPpro")
    #ax.legend()
    graph = FigureCanvasTkAgg(fig, master=FenetreP)
    canvas = graph.get_tk_widget()
    canvas.grid(row=0, column=0)
    TexteP = tk.Tk()
    TexteP.wm_title("Factor de potencia")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1 = tk.Label(TexteP, text = "<FPprom> = {:.3f}".format(moyenneA(Pmoy)) + "  "\
                      + "<FPmax> = {:.3f}".format(moyenneA(Pmax)) + "  "\
                      + "<FPmin> = {:.3f}".format(moyenneA(Pmin)), foreground = "white", background = "#2B00FA")
    label1['font'] = f_label
    label1.pack()
    label2 = tk.Label(TexteP, text= "FPprom mediana = {:.3f}".format(mediane(Pmoy)) + "  " +\
                                    "FPmax mediana = {:.3f}".format(mediane(Pmax)) + "  " +\
                                    "FPmin mediana = {:.3f}".format(mediane(Pmin)),foreground = "white", background = "#2B00FA")                                                                           
    label2['font'] = f_label
    label2.pack()
    label3 = tk.Label(TexteP, text="El factor de potencia maximo encontrado vale {:.3f}".format(max(Pmax)) + "\n"\
                     + "El factor de potencia minimo encontrado vale {:.3f}".format(min(Pmin)),foreground = "white", background = "#2B00FA")
    label3['font'] = f_label
    label3.pack()

def Fp(fichier): #Révisé, actualisé via le .format
    L = extractiondonne(fichier)    
    Pmoy = []
    Pmax = []
    Pmin = []
    NB = []
    N95 = 0
    N97 = 0
    for i in range(1,len(L)): # Contient les indices des différentes lignes    
        NB.append(i)    
    for i in range(1,len(L)):
        Pmoy.append(L[i][89])
        if L[i][89] >= 0.97:
            N95+= 1
            N97+= 1
        if L[i][89] >= 0.95 and L[i][89] < 0.97:
            N95+= 1
    for i in range(1,len(L)):
        Pmax.append(L[i][88])
    for i in range(1,len(L)):
        Pmin.append(L[i][90])
    plt.plot(NB,Pmoy, label = "Factor de potencia")       
    plt.title("Factor de potencia")
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Factor de potencia.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Factor de potencia.png")
        plt.clf()
    document.add_heading("Factor de potencia: ",1)
    document.add_paragraph("Acordamos que el codigo de red menciona que, por ahora, el factor de potencia debe "\
                           + "superar 0.95 sobre 95% del tiempo durante un peridode mensual "\
                           + "con medidas cinco minutal, y a partir del ano 2026, debera superar"\
                           + " el valor de 0.97 sobre 97% del tiempo." + "\n")
    document.add_paragraph("Este grafico mostra la evolución del factor de potentia:" + "\n")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Factor de potencia.png", width=Inches(4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio del factor de potencia vale {:.3f}".format(moyenneA(Pmoy)))
    document.add_paragraph("- El valor promedio del factor de potencia maximo vale {:.3f}".format(moyenneA(Pmax)))                       
    document.add_paragraph("- El valor promedio del factor de potencia maximo vale {:.3f}".format(moyenneA(Pmin)))
    document.add_paragraph("- La mediana del factor de potencia vale {:.3f}".format(mediane(Pmoy)))
    document.add_paragraph("- La mediana del factor de potencia maximo vale {:.3f}".format(mediane(Pmax)))                       
    document.add_paragraph("- La mediana del factor de potencia minimo vale {:.3f}".format(mediane(Pmin)))
    document.add_paragraph("- El factor de potencia maximo encontrado es de {:.3f}".format(max(Pmax)))
    document.add_paragraph("- El factor de potencia minimo encontrado es de {:.3f}".format(min(Pmin)) + "\n")
    if N95 >= 0.95*len(Pmoy) and N97 < 0.97*len(Pmoy):
        document.add_paragraph("El factor de potencia sube 0.95 por {:1f}".format((N95/len(Pmoy))*100) + " % del tiempo"\
                               + "y entonces cumple el codigo de red hasta 2026.")    
    if N97 >= 0.97*len(Pmoy):
        document.add_paragraph("El factor de potencia sube 0.97 por {:1f}".format((N97/len(Pmoy))*100)\
                               + " % del tiempo, y entonces cumple las especificaciones del codigo de red.")
    if N95 <= 0.95*len(Pmoy):
        document.add_paragraph("El factor de potencia no sube 0.95 por 95% del tiempo, y entonces no"\
                               + " cumple las especificaciones del codigo de red." + "\n")
    document.add_paragraph("Además, se aplica un cargo a la facturación debido al factor de potencia"\
                           + "que se define así: " + "\n")
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\Cargo.png", width=Inches(2.8))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("\n" + "Donde el factor de cargo se puede definir de dos maneras, dependiendo del valor del factor de potencia:" + "\n")
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\Factor de cargo.png", width=Inches(3.1))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Así podemos notar que si el factor de potencia no sube los 90%, el cargo es negativo y se vuelve así en un recargo."\
                           + " Al contrario, si el factor de potencia sube los 90%, el cargo es positivo y se vuelve en una bonificatión" + "\n")
    if mediane(Pmoy) > 0.9:
        document.add_paragraph("En nuestro caso, el factor de cargo sera una bonificación y valdra {:.3f}".format(0.25*(1 - 0.9/mediane(Pmoy))))
    if mediane(Pmoy) <= 0.9:
        document.add_paragraph("En nuestro caso, el factor de cargo sera un recargo y valdra {:.3f}".format(0.6*(0.9/mediane(Pmoy) - 1)))


# En una ventana tkinter, voltea al valor promedio de Spro, Smin y Smax, a sus valores medianas, al los valores
# maxima y minima encontradas y en otra ventana, desplega al grafico de Ppro, Pmin y Pmax, y a la media mobil 
# sobre 6 horas         
def BoutonFreq(fichier):  # Révisé !!! #A améliorer éventuellement
    L = extractiondonne(fichier) 
    Pmoy = []
    Pmax = []
    Pmin = []
    NB = []
    for i in range(1,len(L)): # Contient les indices des différentes lignes    
        NB.append(i)
    for i in range(1,len(L)): # LLena una lista con los valores de la potencia activa promedia
        Pmoy.append(L[i][92])    
    for i in range(1,len(L)): # Llena una lista con los valores de la potencia activa maxima
        Pmax.append(L[i][91])
    for i in range(1,len(L)): # Llena una lista con los valores de la potencia activa minima
        Pmin.append(L[i][93])
    FenetreP = tk.Tk()
    FenetreP.wm_title("Frecuencia (Hz)")
    FenetreP.configure(background="#2B00FA")
    fig = Figure(figsize=(6, 4), dpi=96)
    ax = fig.add_subplot(111)
    ax.plot(NB,Pmoy, label = "Fpro")
    if max(Pmoy) >= 60.9:
        ax.axhline(y=61,color='red')
    if min(Pmoy) <= 59.1:
        ax.axhline(y=59,color='red')     
    graph = FigureCanvasTkAgg(fig, master=FenetreP)
    canvas = graph.get_tk_widget()
    canvas.grid(row=0, column=0)
    TexteP = tk.Tk()
    TexteP.wm_title("Frecuencia (Hz)")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1 = tk.Label(TexteP, text = "<Fprom> = {:.3f}".format(moyenneA(Pmoy)) + " Hz" + "  "\
                      + "<Fmax> =  {:.3f}".format(moyenneA(Pmax)) + " Hz" + "  "\
                      + "<Fmin> = {:.3f}".format(moyenneA(Pmin)) + " Hz",foreground = "white", background = "#2B00FA")
    label1['font'] = f_label
    label1.pack()
    label2 = tk.Label(TexteP, text= "Fprom mediana = {:.3f}".format(mediane(Pmoy)) + " Hz" + "  " +\
                                    "Fmax mediana = {:.3f}".format(mediane(Pmax)) + " Hz" + "  " +\
                                    "Fmin mediana = {:.3f}".format(mediane(Pmin)) + " Hz",foreground = "white", background = "#2B00FA")                                                                           
    label2['font'] = f_label
    label2.pack()
    label3 = tk.Label(TexteP, text="La frecuencia maxima encontrada vale {:.3f}".format(max(Pmax)) + " Hz" + "\n"\
                     + "La frecuencia minima encontrada vale {:.3f}".format(min(Pmin)) + " Hz",foreground = "white", background = "#2B00FA")
    label3['font'] = f_label
    label3.pack()

def Freq(fichier): # Révisé !!! #A modifier éventuellement
    L = extractiondonne(fichier)    
    Pmoy = []
    Pmax = []
    Pmin = []
    NB = []
    Mauvais = []
    Bon = 0
    for i in range(1,len(L)): # Contient les indices des différentes lignes    
        NB.append(i)    
    for i in range(1,len(L)):
        Pmoy.append(L[i][92])
        if L[i][92] >= 59 and L[i][92] <= 61:
            Bon+=1
        else:
            Mauvais.append((i,L[i][92]))
    for i in range(1,len(L)):
        Pmax.append(L[i][91])
    for i in range(1,len(L)):
        Pmin.append(L[i][93])
    plt.plot(NB,Pmoy, label = "Frecuencia")
    if max(Pmoy) >= 60.9:
        plt.axhline(y=61,color='red')
    if min(Pmoy) <= 59.1:
        plt.axhline(y=59,color='red')             
    plt.title("Frecuencia (Hz)")
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Frecuencia.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Frecuencia.png")
        plt.clf()
    document.add_heading("Partida 2: Frecuencia ",1)
    document.add_paragraph("Las especificationes CFE L0000-45 y CFE L0000-70 escpecifican que la frecuencia promedia debe de ser de 60 Hz ±"\
                           " 1 Hz por 99,5 % del tiempo." + "\n")
    document.add_paragraph("Este grafico mostra la evolución de la frecuencia:")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Frecuencia.png", width=Inches(4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio de la frecuencia promedia vale {:.3f}".format(moyenneA(Pmoy)) + " Hz.")
    document.add_paragraph("- El valor promedio de la frecuencia maxima vale {:.3f}".format(moyenneA(Pmax)) + " Hz.")                       
    document.add_paragraph("- El valor promedio de la frecuencia minima vale {:.3f}".format(moyenneA(Pmin)) + " Hz.")
    document.add_paragraph("- La mediana de la frecuencia promedia vale {:.3f}".format(mediane(Pmoy)) + " Hz.")
    document.add_paragraph("- La mediana de la frecuencia maxima vale {:.3f}".format(mediane(Pmax)) + " Hz.")                       
    document.add_paragraph("- La mediana de la frecuencia minima vale {:.3f}".format(mediane(Pmin)) + " Hz.")
    document.add_paragraph("- La frecuencia maxima encontrada es de {:.3f}".format(max(Pmax)) + " Hz.")
    document.add_paragraph("- La frecuencia minima encontrada es de {:.3f}".format(min(Pmin)) + " Hz." + "\n")
    if Bon >= 0.995*len(Pmoy):
        document.add_paragraph("La frecuencia esta adentro de la zona de tolerancia por {:.1f}".format((Bon/len(Pmoy))*100)\
                               + " % del tiempo, y asi respecta la norma.")    
    if Bon < 0.995*len(Pmoy):
        document.add_paragraph("La frecuencia esta adentro de la zona de tolerancia por {:.1f}".format((Bon/len(Pmoy))*100)\
                               + " % del tiempo, y entonces no respeta la norma.")
def BoutonTF(fichier): # Révisé !!! A améliorer éventuellement
    L = extractiondonne(fichier)
    Amoy = [] #2
    Bmoy = [] #5
    Cmoy = [] #8
    Dmoy = [] #11
    Amax = L[1][1] #1
    Bmax = L[1][4] #4
    Cmax = L[1][7] #7
    Dmax = L[1][10] #10
    Amin = L[1][3] #3
    Bmin = L[1][6] #6
    Cmin = L[1][9] #9
    Dmin = L[1][12] #12
    NB = []
    for i in range(1,len(L)):
        NB.append(i)
    for i in range(1,len(L)):
        Amoy.append(L[i][2])
        Bmoy.append(L[i][5])
        Cmoy.append(L[i][8])
        Dmoy.append(L[i][11])
        if Amax < L[i][1]:
            Amax = L[i][1]
        if Bmax < L[i][4]:
            Bmax = L[i][4]
        if Cmax < L[i][7]:
            Cmax = L[i][7]
        if Dmax < L[i][10]:
            Dmax = L[i][10]
        if Amin > L[i][3]:
            Amin = L[i][3]
        if Bmin > L[i][6]:
            Bmin = L[i][6]
        if Cmin > L[i][9]:
            Cmin = L[i][9]
        if Dmin > L[i][12]:
            Dmin = L[i][12]
    FenetreP = tk.Tk()
    FenetreP.wm_title("Tension en linea (V)")
    FenetreP.configure(background="#2B00FA")
    fig = Figure(figsize=(6, 4), dpi=96)
    ax = fig.add_subplot(111)
    ax.plot(NB,Amoy, label = "fase A")
    ax.legend()
    ax.plot(NB,Bmoy, label = "fase B")
    ax.legend()
    ax.plot(NB,Cmoy, label = "fase C")
    ax.legend()
    graph = FigureCanvasTkAgg(fig, master=FenetreP)
    canvas = graph.get_tk_widget()
    canvas.grid(row=0, column=0)
    TexteP = tk.Tk()
    TexteP.wm_title("Tension en linea (V)")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1A = tk.Label(TexteP, text = "<Vprom>(fase A) = " + str(ceil(moyenneA(Amoy))) + " V" + "  "\
                      + "Vmax(fase A) =  "\
                     + str(ceil(Amax)) + " V" + "  " +\
                         "Vmin(fase A) = "\
                     + str(ceil(Amin)) + " V" + "  "+\
                         "Vprom mediana (fase A) = " + str(ceil(mediane(Amoy))) + " V", foreground = "white", background = "#2B00FA")
    label1A['font'] = f_label
    label1A.pack()
    label1B = tk.Label(TexteP, text = "<Vprom>(fase B) = " + str(ceil(moyenneA(Bmoy))) + " V" + "  "\
                      + "Vmax(fase B) =  "\
                     + str(ceil(Bmax)) + " V" + "  " +\
                         "Vmin(fase B) = "\
                     + str(ceil(Bmin)) + " V" + "  "+\
                         "Vprom mediana (fase B) = " + str(ceil(mediane(Bmoy))) + " V", foreground = "white", background = "#2B00FA")
    label1B['font'] = f_label
    label1B.pack()
    label1C = tk.Label(TexteP, text = "<Vprom>(fase C) = " + str(ceil(moyenneA(Cmoy))) + " V" + "  "\
                      + "Vmax(fase C) =  "\
                     + str(ceil(Cmax)) + " V" + "  " +\
                         "Vmin(fase C) = "\
                     + str(ceil(Cmin)) + " V" + "  " +\
                         "Vprom mediana (fase C) = " + str(ceil(mediane(Cmoy))) + " V", foreground = "white", background = "#2B00FA")
    label1C['font'] = f_label
    label1C.pack()
    label1D = tk.Label(TexteP, text = "<Vprom> (neutro) = " + str(ceil(moyenneA(Dmoy))) + " V" + "  "\
                      + "Vmax(neutro) =  "\
                     + str(ceil(Dmax)) + " V" + "  " +\
                         "Vmin(neutro) = "\
                     + str(ceil(Dmin)) + " V" + "  "+\
                         "Vprom mediana (neutro)" + str(ceil(mediane(Dmoy))) + " V", foreground = "white", background = "#2B00FA")
    label1D['font'] = f_label
    label1D.pack()                            
    

def TensionF(fichier): # Révisé !!! # A améliorer éventuellement
    L = extractiondonne(fichier)
    Amoy = [] #2
    Bmoy = [] #5
    Cmoy = [] #8
    Dmoy = [] #11
    Amax = L[1][1] #1
    Bmax = L[1][4] #4
    Cmax = L[1][7] #7
    Dmax = L[1][10] #10
    Amin = L[1][3] #3
    Bmin = L[1][6] #6
    Cmin = L[1][9] #9
    Dmin = L[1][12] #12
    NB = []
    for i in range(1,len(L)):
        NB.append(i)
    for i in range(1,len(L)):
        Amoy.append(L[i][2])
        Bmoy.append(L[i][5])
        Cmoy.append(L[i][8])
        Dmoy.append(L[i][11])
        if Amax < L[i][1]:
            Amax = L[i][1]
        if Bmax < L[i][4]:
            Bmax = L[i][4]
        if Cmax < L[i][7]:
            Cmax = L[i][7]
        if Dmax < L[i][10]:
            Dmax = L[i][10]
        if Amin > L[i][3]:
            Amin = L[i][3]
        if Bmin > L[i][6]:
            Bmin = L[i][6]
        if Cmin > L[i][9]:
            Cmin = L[i][9]
        if Dmin > L[i][12]:
            Dmin = L[i][12]
    plt.plot(NB,Amoy, label = r"$V_{A_{pro}}$")
    plt.legend()
    plt.plot(NB,Bmoy, label = r"$V_{B_{pro}}$")
    plt.legend()
    plt.plot(NB,Cmoy, label = r"$V_{C_{pro}}$")
    plt.legend()        
    plt.title("Tension de fase (V)")
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Tension de fase.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Tension de fase.png")
        plt.clf()
        plt.plot(NB,Dmoy, label = "VDmoy")
        #plt.legend()
        plt.title("Tension en el neutro (V)")
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Tension en neutro.png")
        plt.clf()
    document.add_heading("Partida 3: Tensiones y corrientes",1)
    document.add_heading("3.1 Tension de fase: ",2)
    document.add_paragraph("Llamamos tension de fase, la diferencia de potencial entre una fase y el neutro." + "\n")
    document.add_paragraph("Este grafico mostra la evolucion de la tension de fase por las tres fases:" + "\n")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Tension de fase.png", width=Inches(5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio de la tension en la fase A vale {:.1f}".format(moyenneA(Amoy)) + " V.")
    document.add_paragraph("- El valor promedio de la tension en la fase B vale {:.1f}".format(moyenneA(Bmoy)) + " V.")                       
    document.add_paragraph("- El valor promedio de la tension en la fase C vale {:.1f}".format(moyenneA(Cmoy)) + " V.")
    document.add_paragraph("- La mediana de la tension en la fase A vale {:.1f}".format(mediane(Amoy)) + " V.")
    document.add_paragraph("- La mediana de la tension en la fase B vale {:.1f}".format(mediane(Bmoy)) + " V.")                       
    document.add_paragraph("- La mediana de la tension en la fase C vale {:.1f}".format(mediane(Cmoy)) + " V.")
    document.add_paragraph("- La potencia maxima encontrada en la fase A vale {:.1f}".format(Amax) + " V.")
    document.add_paragraph("- La potencia maxima encontrada en la fase B vale {:.1f}".format(Bmax) + " V.")
    document.add_paragraph("- La potencia maxima encontrada en la fase C vale {:.1f}".format(Cmax) + " V.")
    document.add_paragraph("- La potencia minima encontrada en la fase A vale {:.1f}".format(Amin) + " V.")
    document.add_paragraph("- La potencia minima encontrada en la fase B vale {:.1f}".format(Bmin) + " V.")
    document.add_paragraph("- La potencia minima encontrada en la fase C vale {:.1f}".format(Cmin) + " V." + "\n")
    document.add_paragraph(" Igualmente, este grafico nos mostra la tension entre el neutro y la tierra: " + "\n")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Tension en neutro.png", width=Inches(5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos tambien observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio de la tension en el neutro vale {:.1f}".format(moyenneA(Dmoy)) + " V.")
    document.add_paragraph("- La mediana de la tension en el neutro vale {:.1f}".format(mediane(Dmoy)) + " V.")
    document.add_paragraph("- La potencia maxima encontrada en el neutro vale {:.1f}".format(Dmax) + " V.")
    document.add_paragraph("- La potencia minima encontrada en el neutro vale {:.1f}".format(Dmin) + " V." + "\n")

def BoutonTL(fichier): #Révisé !!! #A améliorer
    L = extractiondonne(fichier)
    ABmoy = [] #14
    BCmoy = [] #17
    CAmoy = [] #20
    ABmax = [] #13
    BCmax = [] #16
    CAmax = [] #19
    ABmin = [] #15
    BCmin = []#18
    CAmin = [] #21    
    NB = []
    for i in range(1,len(L)):
        NB.append(i)
    for i in range(1,len(L)):
        ABmax.append(L[i][13])
        ABmoy.append(L[i][14])
        ABmin.append(L[i][15])
        BCmax.append(L[i][16])
        BCmoy.append(L[i][17])
        BCmin.append(L[i][18])
        CAmax.append(L[i][19])
        CAmoy.append(L[i][20])
        CAmin.append(L[i][21])
    FenetreP = tk.Tk()
    FenetreP.wm_title("Tension entre fases (V)")
    FenetreP.configure(background="#2B00FA")
    fig = Figure(figsize=(6, 4), dpi=96)
    ax = fig.add_subplot(111)
    ax.plot(NB,ABmoy, label = "V A-B")
    ax.legend()
    ax.plot(NB,BCmoy, label = "V B-C")
    ax.legend()
    ax.plot(NB,CAmoy, label = "V C-A")
    ax.legend()
    graph = FigureCanvasTkAgg(fig, master=FenetreP)
    canvas = graph.get_tk_widget()
    canvas.grid(row=0, column=0)
    TexteP = tk.Tk()
    TexteP.wm_title("Tension entre fases (V)")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1A = tk.Label(TexteP, text = "<Vprom>(fases A-B) = " + str(ceil(moyenneA(ABmoy))) + " V" + "  "\
                      + "Vmax(fases A-B) =  "\
                     + str(ceil(max(ABmax))) + " V" + "  " +\
                         "Vmin(fases A-B) = "\
                     + str(ceil(min(ABmin))) + " V" + "  "+\
                         "Vprom mediana (fases A-B) = " + str(ceil(mediane(ABmoy))) + " V", foreground = "white", background = "#2B00FA")
    label1A['font'] = f_label
    label1A.pack()
    label1B = tk.Label(TexteP, text = "<Vprom>(fases B-C) = " + str(ceil(moyenneA(BCmoy))) + " V" + "  "\
                      + "Vmax(fases B-C) =  "\
                     + str(ceil(max(BCmax))) + " V" + "  " +\
                         "Vmin(fases B-C) = "\
                     + str(ceil(min(BCmin))) + " V" + "  "+\
                         "Vprom mediana (fases B-C) = " + str(ceil(mediane(BCmoy))) + " V", foreground = "white", background = "#2B00FA")
    label1B['font'] = f_label
    label1B.pack()
    label1C = tk.Label(TexteP, text = "<Vprom>(fases C-A) = " + str(ceil(moyenneA(CAmoy))) + " V" + "  "\
                      + "Vmax(fases C-A) =  "\
                     + str(ceil(max(CAmax))) + " V" + "  " +\
                         "Vmin(fases C-A) = "\
                     + str(ceil(min(CAmin))) + " V" + "  " +\
                         "Vprom mediana (fases C-A) = " + str(ceil(mediane(CAmoy))) + " V", foreground = "white", background = "#2B00FA")
    label1C['font'] = f_label
    label1C.pack() 
    
    
def TensionL(fichier): #Révisé !!! #A améliorer
    L = extractiondonne(fichier)
    ABmoy = [] #14
    BCmoy = [] #17
    CAmoy = [] #20
    ABmax = [] #13
    BCmax = [] #16
    CAmax = [] #19
    ABmin = [] #15
    BCmin = []#18
    CAmin = [] #21    
    NB = []
    for i in range(1,len(L)):
        NB.append(i)
    for i in range(1,len(L)):
        ABmax.append(L[i][13])
        ABmoy.append(L[i][14])
        ABmin.append(L[i][15])
        BCmax.append(L[i][16])
        BCmoy.append(L[i][17])
        BCmin.append(L[i][18])
        CAmax.append(L[i][19])
        CAmoy.append(L[i][20])
        CAmin.append(L[i][21])
    plt.plot(NB,ABmoy, label = r"$V_{{A-B}_{pro}}$")
    plt.legend()
    plt.plot(NB,BCmoy, label = r"$V_{{B-C}_{pro}}$")
    plt.legend()
    plt.plot(NB,CAmoy, label = r"$V_{{C-A}_{pro}}$")
    plt.legend()        
    plt.title("Tension entre fases (V)")
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Tension de linea.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Tension de linea.png")
        plt.clf()
    document.add_heading("3.2 Tension de linea: ",2)
    document.add_paragraph("LLamamos tensión de linea la differencia de potenciales entre las fases." + "\n")                            
    document.add_paragraph("Este grafico mostra la evolucion de la tension de linea entre las tres fases:" + "\n")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Tension de linea.png", width=Inches(4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio de la tension entre fase A y fase B vale {:.1f}".format(moyenneA(ABmoy)) + " V.")
    document.add_paragraph("- El valor promedio de la tension entre fase B y fase C vale {:.1f}".format(moyenneA(BCmoy)) + " V.")                       
    document.add_paragraph("- El valor promedio de la tension entre fase C y fase A vale {:.1f}".format(moyenneA(CAmoy)) + " V.")
    document.add_paragraph("- La mediana de la tension entre fase A y fase B vale {:.1f}".format(mediane(ABmoy)) + " V.")
    document.add_paragraph("- La mediana de la tension entre fase B y fase C vale {:.1f}".format(mediane(BCmoy)) + " V.")                       
    document.add_paragraph("- La mediana de la tension entre fase C y fase A vale {:.1f}".format(mediane(CAmoy)) + " V.")
    document.add_paragraph("- La tension maxima encontrada entre fase A y fase B vale {:.1f}".format(max(ABmax)) + " V.")
    document.add_paragraph("- La tension maxima encontrada entre fase B y fase C vale {:.1f}".format(max(BCmax)) + " V.")
    document.add_paragraph("- La tension maxima encontrada entre fase C y fase A vale {:.1f}".format(max(CAmax)) + " V.")
    document.add_paragraph("- La tension minima encontrada entre fase A y fase B vale {:.1f}".format(min(ABmin)) + " V.")
    document.add_paragraph("- La tension minima encontrada entre fase B y fase C vale {:.1f}".format(min(BCmin)) + " V.")
    document.add_paragraph("- La tension minima encontrada entre fase C y fase A vale {:.1f}".format(min(CAmin)) + " V.")

def BoutonDesT(fichier): 
    L = extractiondonne(fichier)
    Dmoy = []
    Mauvais = []
    NB = []
    Hihi = []
    for i in range(1,len(L)):
        Hihi.append(L[i][2])
    for i in range(1,len(L)):
        NB.append(i)
    for i in range(1,len(L)):
        VAmoy = L[i][2]
        VBmoy = L[i][5]
        VCmoy = L[i][8]
        prom = (1/3)*(VAmoy + VBmoy + VCmoy)
        ecart = abs(max(VAmoy,VBmoy,VCmoy) - prom)
        Dmoy.append((ecart/prom)*100)
        if mediane(Hihi) <= 1000 and (ecart/prom)*100 >= 3:
            Mauvais.append([i,(ecart/prom)*100])
        if mediane(Hihi) > 1000 and (ecart/prom)*100 >= 2:
            Mauvais.append([i,(ecart/prom)*100])
    FenetreP = tk.Tk()
    FenetreP.wm_title("Desbalanceo de voltaje (%)")
    FenetreP.configure(background="#2B00FA")
    fig = Figure(figsize=(6, 4), dpi=96)
    ax = fig.add_subplot(111)
    ax.plot(NB,Dmoy)
    if max(Dmoy) >= 2:
        ax.axhline(y=3,color='red')    
    graph = FigureCanvasTkAgg(fig, master=FenetreP)
    canvas = graph.get_tk_widget()
    canvas.grid(row=0, column=0)
    TexteP = tk.Tk()
    TexteP.wm_title("Desbalanceo de voltaje (%)")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1A = tk.Label(TexteP, text = "<Desbalanceo de voltaje> = {:.3f}".format(moyenneA(Dmoy)) + " %" + "  "\
                      + "max(desbalanceo) =  {:.3f}".format(max(Dmoy)) + " %" + "  "\
                      + "min(desbalanceo) = {:.3f}".format(min(Dmoy))+ " %", foreground = "white", background = "#2B00FA")
    label1A['font'] = f_label
    label1A.pack()        

def DesT(fichier):
    L = extractiondonne(fichier)
    Dmoy = []
    Mauvais = []
    NB = []
    bon = 0
    for i in range(1,len(L)):
        NB.append(i)
    for i in range(1,len(L)):
        VAmoy = L[i][2]
        VBmoy = L[i][5]
        VCmoy = L[i][8]
        prom = (1/3)*(VAmoy + VBmoy + VCmoy)
        ecart = abs(max(VAmoy,VBmoy,VCmoy) - prom)
        Dmoy.append((ecart/prom)*100)
        if TN < 1000 and (ecart/prom)*100 > 3:
            Mauvais.append([i,(ecart/prom)*100])
        if TN >= 1000 and (ecart/prom)*100 > 2:
            Mauvais.append([i,(ecart/prom)*100])
        if TN < 1000 and (ecart/prom)*100 <= 3:
            bon+=1
        if TN >= 1000 and (ecart/prom)*100 <= 2:
            bon+=1
    plt.plot(NB,Dmoy)
    if max(Dmoy) >= 2:
        plt.axhline(y=3,color='red')
    plt.title("Desbalance de tensión (%)")
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Desbalanceo de Voltaje.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Desbalanceo de Voltaje.png")
        plt.clf() 
    document.add_heading("Partida 4: Desbalances ",1)
    document.add_heading("1.1 Desbalance de tensión: ",2)
    document.add_paragraph("Este grafico mostra el desbalance de tensión: " + "\n")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Desbalanceo de Voltaje.png", width=Inches(4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("El desbalanceo promedio vale: {:.1f}".format(moyenneA(Dmoy)) + " %")
    document.add_paragraph("El desbalanceo mediano vale: {:.1f}".format(mediane(Dmoy)) + " %")
    document.add_paragraph("El desbalanceo maximo vale: {:.1f}".format(max(Dmoy)) + " %")
    document.add_paragraph("El desbalanceo minimo vale: {:.1f}".format(min(Dmoy)) + " %" + "\n")
    document.add_paragraph("Recordamos que las especificaciones CFE L0000-45 y CFE L0000-70 especifican que el desbalanceo de tension no supere esos valores por 95% del tiempo: " + "\n")    
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\Tolerance desequilibre en tension.png", width=Inches(4))
    if TN >= 1000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance es inferior a 2% en {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo"\
                               + " y entonces cumple la norma." + "\n")    
    if TN < 1000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance es inferior a 3% en {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo"\
                               + " y entonces cumple la norma." + "\n")
    if TN >= 1000 and bon/len(Dmoy) < 0.95:
        document.add_paragraph("El desbalance es inferior a 2% en {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo"\
                               + " y entonces no cumple la norma." + "\n")
    if TN < 1000 and bon/len(Dmoy) < 0.95:
        document.add_paragraph("El desbalance es inferior a 3% en {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo"\
                               + " y entonces no cumple la norma." + "\n")
        
# Bouton pour déséquilibre de courant
def BoutonDestI(fichier):
    L = extractiondonne(fichier)
    Dmoy = []
    Mauvais = []
    NB = []
    for i in range(1,len(L)):
        NB.append(i)
    for i in range(1,len(L)):
        VAmoy = L[i][23]
        VBmoy = L[i][26]
        VCmoy = L[i][29]
        prom = (1/3)*(VAmoy + VBmoy + VCmoy)
        ecart = abs(max(VAmoy,VBmoy,VCmoy) - prom)
        Dmoy.append((ecart/prom)*100)
    FenetreP = tk.Tk()
    FenetreP.wm_title("Desbalanceo de voltaje (%)")
    FenetreP.configure(background="#2B00FA")
    fig = Figure(figsize=(6, 4), dpi=96)
    ax = fig.add_subplot(111)
    ax.plot(NB,Dmoy)    
    if Icc/Inom < 20 and TN < 1000 and max(Dmoy) >= 4.5:
        ax.axhline(y=5,color='red') 
    if Icc/Inom < 20 and 1000 <= TN < 35000 and max(Dmoy) >= 2:
        ax.axhline(y= 2.5,color='red')
    if Icc/Inom < 20 and TN >= 35000 and max(Dmoy) >= 2:    
        ax.axhline(y=2.5,color='red')
    if 20 <= Icc/Inom < 50 and TN < 1000 and max(Dmoy) >= 7.5:    
        ax.axhline(y=8,color='red')
    if 20 <= Icc/Inom < 50 and 1000 <= TN < 35000 and max(Dmoy) >= 3.5:    
        ax.axhline(y=4,color='red')
    if 20 <= Icc/Inom < 50 and TN >= 35000 and max(Dmoy) >= 2.5:    
        ax.axhline(y=3,color='red')
    if 50 <= Icc/Inom < 100 and TN < 1000 and max(Dmoy) >= 11.5:
        ax.axhline(y=12,color='red')
    if 50 <= Icc/Inom < 100 and 1000 <= TN < 35000 and max(Dmoy) >= 5.5:
        ax.axhline(y=6,color='red')
    if 50 <= Icc/Inom < 100 and TN >= 35000 and max(Dmoy) >=3.3:
        ax.axhline(y=3.75,color='red')
    if 100 <= Icc/Inom < 1000 and TN < 1000 and max(Dmoy) >= 14.5:
        ax.axhline(y = 15,color='red')
    if 100 <= Icc/Inom < 1000 and 1000 <= TN < 35000 and max(Dmoy) >= 7:
        ax.axhline(y=7.5,color='red')
    if 100 <= Icc/Inom < 1000 and TN >= 35000 and max(Dmoy) >= 3.5:
        ax.axhline(y=4,color='red')
    if Icc/Inom >= 1000 and TN < 1000 and max(Dmoy) >= 19.5:
        ax.axhline(y=20,color='red')
    if Icc/Inom >= 1000 and 1000 <= TN < 35000 and (ecart/prom)*100 >= 9.5:
        ax.axhline(y=10,color='red')
    if Icc/Inom >= 1000 and TN >= 35000 and (ecart/prom)*100 >= 4.5:
        ax.axhline(y=5,color='red')
    graph = FigureCanvasTkAgg(fig, master=FenetreP)
    canvas = graph.get_tk_widget()
    canvas.grid(row=0, column=0)
    TexteP = tk.Tk()
    TexteP.wm_title("Desbalanceo de corriente (%)")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1A = tk.Label(TexteP, text = "<Desbalanceo de corriente> = {:.3f}".format(moyenneA(Dmoy)) + " %" + "  "\
                      + "max(desbalanceo) =  {:.3f}".format(max(Dmoy)) + " %" + "  "\
                      + "min(desbalanceo) = {:.3f}".format(min(Dmoy)) + " %",foreground = "white", background = "#2B00FA")
    label1A['font'] = f_label
    label1A.pack()        

# Deséquilibre de courant    
def DestI(fichier):
    L = extractiondonne(fichier)
    Dmoy = []
    Mauvais = []
    NB = []
    R = Icc/Inom
    bon = 0
    for i in range(1,len(L)):
        NB.append(i)
    for i in range(1,len(L)):
        VAmoy = L[i][23]
        VBmoy = L[i][26]
        VCmoy = L[i][29]
        prom = (1/3)*(VAmoy + VBmoy + VCmoy)
        ecart = abs(max(VAmoy,VBmoy,VCmoy) - prom)
        Dmoy.append((ecart/prom)*100)
        if prom != 0:
            if R < 20 and TN < 1000 and (ecart/prom)*100 > 5:
                Mauvais.append([i,(ecart/prom)*100])
            if R < 20 and 1000 <= TN < 35000 and (ecart/prom)*100 > 2.5:
                Mauvais.append([i,(ecart/prom)*100])    
            if R < 20 and TN >= 35000 and (ecart/prom)*100 > 2.5:    
                Mauvais.append([i,(ecart/prom)*100])
            if R < 20 <= Icc/Inom < 50 and TN < 1000 and (ecart/prom)*100 > 8:    
                Mauvais.append([i,(ecart/prom)*100])          
            if 20 <= R < 50 and 1000 <= TN < 35000 and (ecart/prom)*100 > 4:    
                Mauvais.append([i,(ecart/prom)*100])
            if 20 <= R < 50 and TN >= 35000 and (ecart/prom)*100 > 3:    
                Mauvais.append([i,(ecart/prom)*100])
            if 50 <= R < 100 and TN < 1000 and (ecart/prom)*100 > 12:
                Mauvais.append([i,(ecart/prom)*100])
            if 50 <= R < 100 and 1000 <= TN < 35000 and (ecart/prom)*100 > 6:
                Mauvais.append([i,(ecart/prom)*100])
            if 50 <= R < 100 and TN >= 35000 and (ecart/prom)*100 > 3.75:
                Mauvais.append([i,(ecart/prom)*100])
            if 100 <= R < 1000 and TN < 1000 and (ecart/prom)*100 > 15:
                Mauvais.append([i,(ecart/prom)*100])
            if 100 <= R < 1000 and 1000 <= TN < 35000 and (ecart/prom)*100 > 7.5:
                Mauvais.append([i,(ecart/prom)*100])
            if 100 <= R < 1000 and TN >= 35000 and (ecart/prom)*100 >= 4:
                Mauvais.append([i,(ecart/prom)*100])
            if R >= 1000 and TN < 1000 and (ecart/prom)*100 >= 20:
                Mauvais.append([i,(ecart/prom)*100])
            if R >= 1000 and 1000 <= TN < 35000 and (ecart/prom)*100 >= 10:
                Mauvais.append([i,(ecart/prom)*100])
            if R >= 1000 and TN >= 35000 and (ecart/prom)*100 >= 5:
                Mauvais.append([i,(ecart/prom)*100])            
    plt.plot(NB,Dmoy)
    if Icc/Inom < 20 and TN < 1000 and max(Dmoy) >= 4.5:
        plt.axhline(y=5,color='red') 
    if Icc/Inom < 20 and 1000 <= TN < 35000 and max(Dmoy) >= 2:
        plt.axhline(y= 2.5,color='red')
    if Icc/Inom < 20 and TN >= 35000 and max(Dmoy) >= 2:    
        plt.axhline(y=2.5,color='red')
    if 20 <= Icc/Inom < 50 and TN < 1000 and max(Dmoy) >= 7.5:    
        plt.axhline(y=8,color='red')
    if 20 <= Icc/Inom < 50 and 1000 <= TN < 35000 and max(Dmoy) >= 3.5:    
        plt.axhline(y=4,color='red')
    if 20 <= Icc/Inom < 50 and TN >= 35000 and max(Dmoy) >= 2.5:    
        plt.axhline(y=3,color='red')
    if 50 <= Icc/Inom < 100 and TN < 1000 and max(Dmoy) >= 11.5:
        plt.axhline(y=12,color='red')
    if 50 <= Icc/Inom < 100 and 1000 <= TN < 35000 and max(Dmoy) >= 5.5:
        plt.axhline(y=6,color='red')
    if 50 <= Icc/Inom < 100 and TN >= 35000 and max(Dmoy) >=3.3:
        plt.axhline(y=3.75,color='red')
    if 100 <= Icc/Inom < 1000 and TN < 1000 and max(Dmoy) >= 14.5:
        plt.axhline(y = 15,color='red')
    if 100 <= Icc/Inom < 1000 and 1000 <= TN < 35000 and max(Dmoy) >= 7:
        plt.axhline(y=7.5,color='red')
    if 100 <= Icc/Inom < 1000 and TN >= 35000 and max(Dmoy) >= 3.5:
        plt.axhline(y=4,color='red')
    if Icc/Inom >= 1000 and TN < 1000 and max(Dmoy) >= 19.5:
        plt.axhline(y=20,color='red')
    if Icc/Inom >= 1000 and 1000 <= TN < 35000 and (ecart/prom)*100 >= 9.5:
        plt.axhline(y=10,color='red')
    if Icc/Inom >= 1000 and TN >= 35000 and (ecart/prom)*100 >= 4.5:
        plt.axhline(y=5,color='red')
    plt.title("Desbalance de corriente (%)")
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Desbalanceo de Corriente.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Desbalanceo de Corriente.png")
        plt.clf() 
    bon = len(Dmoy) - len(Mauvais)
    document.add_heading("4.2 Desbalance de corriente: ",2)
    document.add_paragraph("Este grafico mostra el desbalance de corriente: " + "\n")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Desbalanceo de Corriente.png", width=Inches(5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("El desbalanceo promedio vale: {:.2f}".format(moyenneA(Dmoy)) + " %.")
    document.add_paragraph("El desbalanceo mediano vale: {:.2f}".format(moyenneA(Dmoy)) + " %.")
    document.add_paragraph("El desbalanceo maximo vale: {:.2f}".format(max(Dmoy)) + " %.")
    document.add_paragraph("El desbalanceo minimo vale: {:.2f}".format(min(Dmoy)) + " %." + "\n")
    document.add_paragraph("Recordamos que las especificaciones CFE L0000-45 y CFE L0000-70 especifican que el desbalance de corriente no supere esos valores por 95% del tiempo: " + "\n")    
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\Tolerance desequilibre en courant.png", width=Inches(4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("En nuestro caso, Icc/Il = {:.1f}".format(Icc/Inom) + " y la tension nominal vale {:.0f}".format(TN) + " V." + "\n")
    if Icc/Inom < 20 and TN < 1000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 5% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")    
    if Icc/Inom < 20 and TN < 1000 and bon/len(Dmoy) < 0.95:
        document.add_paragraph("El desbalance esta abajo de 5% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")    
    if Icc/Inom < 20 and 1000 <= TN < 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 2.5% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if Icc/Inom < 20 and 1000 <= TN < 35000 and bon/len(Dmoy) < 0.95:
        document.add_paragraph("El desbalance esta abajo de 2.5% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if Icc/Inom < 20 and TN >= 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 2.5% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if Icc/Inom < 20 and TN >= 35000 and bon/len(Dmoy) < 0.95:
        document.add_paragraph("El desbalance esta abajo de 2.5% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if 20 <= Icc/Inom < 50  and TN < 1000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 8% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if 20 <= Icc/Inom < 50  and TN < 1000 and bon/len(Dmoy) < 0.95:
        document.add_paragraph("El desbalance esta abajo de 8% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if 20 <= Icc/Inom < 50  and 1000 <= TN < 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 4% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if 20 <= Icc/Inom < 50  and 1000 <= TN < 35000 and bon/len(Dmoy) < 0.95:
        document.add_paragraph("El desbalance esta abajo de 4% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if 20 <= Icc/Inom < 50  and  TN >= 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 3% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if 20 <= Icc/Inom < 50  and  TN >= 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 3% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if 50 <= Icc/Inom < 100  and TN < 1000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 12% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if 50 <= Icc/Inom < 100  and TN < 1000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 12% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if 50 <= Icc/Inom < 100  and 1000 <= TN < 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 6% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if 50 <= Icc/Inom < 100  and 1000 <= TN < 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 6% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if 50 <= Icc/Inom < 100  and TN >= 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 3,75% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if 50 <= Icc/Inom < 100  and TN >= 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 3,75% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if  100 <= Icc/Inom < 1000  and TN < 1000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 15% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if  100 <= Icc/Inom < 1000  and TN < 1000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 15% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if 100 <= Icc/Inom < 1000  and 1000 <= TN < 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 7,5% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if 100 <= Icc/Inom < 1000  and 1000 <= TN < 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 7,5% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if 100 <= Icc/Inom < 1000  and TN >= 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 4% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if 100 <= Icc/Inom < 1000  and TN >= 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 4% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if Icc/Inom >= 1000  and TN < 1000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalanceo esta abajo de 20% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if Icc/Inom >= 1000  and TN < 1000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalanceo esta abajo de 20% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if Icc/Inom >= 1000  and 1000 <= TN  < 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 10% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")
    if Icc/Inom >= 1000  and 1000 <= TN  < 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 10% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")
    if Icc/Inom >= 1000  and TN >= 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 5% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces respeta la norma.")    
    if Icc/Inom >= 1000  and TN >= 35000 and bon/len(Dmoy) >= 0.95:
        document.add_paragraph("El desbalance esta abajo de 5% por {:.1f}".format(100*bon/len(Dmoy)) + " % del tiempo y entonces no respeta la norma.")    
    
def BoutonDistT(fichier):
    L = extractiondonne(fichier)
    NB = []
    H = []
    CAIMT = 0
    DATT = 0
    A = []
    Reste = []
    HarmoMax = 0
    for i in range(1,len(L)):
        NB.append(i)
    for i in range (106,106+3*49+1,3):
        for j in range(1,len(L)):
            A.append(L[j][i])
        H.append(moyenneA(A))
        A = []
    CC = H[0]
    for i in range(2,50):
        Reste.append(H[i])
        DATT += (H[i]**2)/(H[1]**2)
    HarmoMax = max(CC,max(Reste))
    CAIMT = 100*(HarmoMax/H[1])
    DATT = 100*sqrt(DATT)
    TexteP = tk.Tk()
    TexteP.wm_title("Distorsión de tensión (%)")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1A = tk.Label(TexteP, text = "CAIMT = {:.2f}".format(CAIMT) + " %" + "  "\
                      + "DATT =  {:.2f}".format(DATT) + " %",foreground = "white", background = "#2B00FA")        
    label1A['font'] = f_label
    label1A.pack()        
    if TN < 1000 and CAIMT <= 6 and DATT <= 8:
        label2A = tk.Label(TexteP, text = "CAIMT y DATT respetan la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()            
    if TN < 1000 and CAIMT <= 6 and DATT > 8:
        label2A = tk.Label(TexteP, text = "DATT no respeta la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()
    if TN < 1000 and CAIMT > 6 and DATT <= 8:
        label2A = tk.Label(TexteP, text = "CAIMT no respeta la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()    
    if TN < 1000 and CAIMT > 6 and DATT > 8:
        label2A = tk.Label(TexteP, text = "CAIMT y DATT no respetan la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()
    if 1000 <= TN < 35000 and CAIMT <= 5 and DATT <= 6.5:        
        label2A = tk.Label(TexteP, text = "CAIMT y DATT respetan la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()
    if 1000 <= TN < 35000 and CAIMT <= 5 and DATT > 6.5:
        label2A = tk.Label(TexteP, text = "DATT no respeta la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()
    if 1000 <= TN < 35000 and CAIMT > 5 and DATT <= 6.5:
        label2A = tk.Label(TexteP, text = "CAIMT no respeta la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()     
    if 1000 <= TN < 35000 and CAIMT > 5 and DATT > 6.5:
        label2A = tk.Label(TexteP, text = "CAIMT y DATT no respetan la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()
    if TN >= 35000 and CAIMT <= 2 and DATT <= 3:
        label2A = tk.Label(TexteP, text = "CAIMT y DATT respetan la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()
    if TN >= 35000 and CAIMT <= 2 and DATT > 3:
        label2A = tk.Label(TexteP, text = "DATT no respeta la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()
    if TN >= 35000 and CAIMT > 2 and DATT <= 3:
        label2A = tk.Label(TexteP, text = "CAIMT no respeta la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()    
    if TN >= 35000 and CAIMT > 2 and DATT > 3:
        label2A = tk.Label(TexteP, text = "CAIMT y DATT no respetan la especificacion",foreground = "white", background = "#2B00FA")        
        label2A['font'] = f_label
        label2A.pack()

def DistT(fichier):
    L = extractiondonne(fichier)
    NB = []
    H = []
    CAIMT = 0
    DATT = 0
    A = []
    Reste = []
    hihi = []
    HarmoMax = 0
    for i in range(1,len(L)):
        NB.append(i)
    for i in range (106,106+3*49+1,3):
        for j in range(1,len(L)):
            A.append(L[j][i])
        H.append(mediane(A))
        A = []
    CC = H[0]
    for i in range(2,50):
        Reste.append(H[i])
        DATT += (H[i]**2)/(H[1]**2)
    HarmoMax = max(CC,max(Reste))
    CAIMT = 100*(HarmoMax/H[1])
    DATT = 100*sqrt(DATT)
    document.add_heading("Distorsión de tensión: " + "\n",1)
    document.add_paragraph("Recordamos que la especificacion CFE-L000045 define dos valores fondamentales"\
                           + " para evaluar la distorsión de tension: el CAIMT y el DATT. Esos valores"\
                           + " son definidos asi:" + "\n")
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\CAIMTDATT.png", width=Inches(4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("En nuestro caso, CAIMT = {:.1f}".format(CAIMT) + " % y DATT = {:.1f}".format(DATT) + " %" + "\n")
    document.add_paragraph("La especificacion CFE-L000045 especifica las siguientes tolerancias sobre esos valores: " + "\n")    
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\Tolérance distortion en tension.png", width=Inches(4))
    if TN < 1000 and CAIMT <= 6 and DATT <= 8:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " la distorcion de tension si respeta la especificacion.")
    if TN < 1000 and CAIMT <= 6 and DATT > 8:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " podemos observar que el CAIMT respecta la especificacion, pero"\
                               + " que no es el caso del DATT.")
    if TN < 1000 and CAIMT > 6 and DATT <= 8:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " podemos observar que el DATT respecta la especificacion, pero"\
                               + " que no es el caso del CAIMT.")    
    if TN < 1000 and CAIMT > 6 and DATT > 8:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " podemos observar que el CAIMT y el DATT no respetan la especificacion.")
    if 1000 <= TN < 35000 and CAIMT <= 5 and DATT <= 6.5:        
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " la distorcion de tension si respeta la especificacion.")
    if 1000 <= TN < 35000 and CAIMT <= 5 and DATT > 6.5:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " podemos observar que el CAIMT respecta la especificacion, pero"\
                               + " que no es el caso del DATT.")
    if 1000 <= TN < 35000 and CAIMT > 5 and DATT <= 6.5:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " podemos observar que el DATT respecta la especificacion, pero"\
                               + " que no es el caso del CAIMT.")        
    if 1000 <= TN < 35000 and CAIMT > 5 and DATT > 6.5:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " podemos observar que el CAIMT y el DATT no respetan la especificacion.")
    if TN >= 35000 and CAIMT <= 2 and DATT <= 3:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " la distorcion de tension si respeta la especificacion.")
    if TN >= 35000 and CAIMT <= 2 and DATT > 3:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " podemos observar que el CAIMT respecta la especificacion, pero"\
                               + " que no es el caso del DATT.")
    if TN >= 35000 and CAIMT > 2 and DATT <= 3:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " podemos observar que el DATT respecta la especificacion, pero"\
                               + " que no es el caso del CAIMT.")    
    if TN >= 35000 and CAIMT > 2 and DATT > 3:
        document.add_paragraph("Dando que la tension nominal vale aqui {:.0f}".format(TN) + " V"\
                               + " podemos observar que el CAIMT y el DATT no respetan la especificacion. " + "\n")
    for i in range(len(H)):
        if TN < 1000 and H[i]/H[1]*100 <= 6 and i!=1:
            hihi.append([str(i),(H[i]/H[1])*100,"sí"])
        if TN < 1000 and H[i]/H[1]*100 > 6 and i!=1:
            hihi.append([str(i),(H[i]/H[1])*100,"no"])
        if  i==1:
            hihi.append([str(i),(H[i]/H[1])*100,"X"])
        if 1000 <= TN < 35000 and H[i]/H[1]*100 <= 5 and i!=1:
            hihi.append([str(i),H[i]/H[1]*100,"sí"])
        if 1000 <= TN < 35000  and H[i]/H[1]*100 > 5 and i!=1:
            hihi.append([str(i),H[i]/H[1]*100,"no"])
        if TN >= 35000 and H[i]/H[1]*100 <= 2 and i!=1:
            hihi.append([str(i),H[i]/H[1]*100,"sí"])
        if TN >= 35000 and H[i]/H[1]*100 > 2 and i!=1:
            hihi.append([str(i),H[i]/H[1]*100,"no"])        
    document.add_paragraph("Para más informaciones sobre las armónicas de tensión, podemos referirse a la siguiente tabla: " + "\n")
    menuTable = document.add_table(rows=1,cols=3)
    #menuTable.style = "Table grid"
    hdr_Cells = menuTable.rows[0].cells
    hdr_Cells[0].text = "número"
    hdr_Cells[1].text = "Vh/V1(%)"
    hdr_Cells[2].text = "¿cumple?"
    for a,b,c in hihi:
        row_Cells = menuTable.add_row().cells
        row_Cells[0].text = a
        row_Cells[1].text = "{:.2f}".format(b)
        row_Cells[2].text = c

def BoutonDistI(fichier):
    L = extractiondonne(fichier)
    NB = []
    H = []
    DATD = 0
    A = []
    R = Icc/Inom
    for i in range(1,len(L)):
        NB.append(i)
    for i in range (106+3*49+3,106+3*49+3 + 3*49 + 1,3):
        for j in range(1,len(L)):
            A.append(L[j][i])
        H.append(mediane(A))
        A = []
    HP11 = [H[0],H[2],H[4],H[6],H[8],H[10]]
    HI11 = [H[3],H[5],H[7],H[9]]
    HP1117 = [H[12],H[14],H[16]]
    HI1117 = [H[11],H[13],H[15]]
    HP1723 = [H[18],H[20],H[22]]
    HI1723 = [H[17],H[19],H[21]]
    HP2335 = [H[24],H[26],H[28],H[30],H[32],H[34]]
    HI2335 = [H[23],H[25],H[27],H[29],H[31],H[33]]
    HP35 = [H[36],H[38],H[40],H[42],H[44],H[46],H[48]]
    HI35 = [H[35],H[37],H[39],H[41],H[43],H[45],H[47]]    
    for i in range(2,50):
        DATD += (H[i]**2)/(H[1]**2)
    DATD = 100*sqrt(DATD)
    CAIMCP11 = 100*(max(HP11)/H[1])
    CAIMCI11 = 100*(max(HI11)/H[1])
    CAIMCP1117 = 100*(max(HP1117)/H[1])
    CAIMCI1117 = 100*(max(HI1117)/H[1])
    CAIMCP1723 = 100*(max(HP1723)/H[1])
    CAIMCI1723 = 100*(max(HI1723)/H[1])
    CAIMCP2335 = 100*(max(HP2335)/H[1])
    CAIMCI2335 = 100*(max(HI2335)/H[1])
    CAIMCP35 = 100*(max(HP35)/H[1])
    CAIMCI35 = 100*(max(HI35)/H[1])
    TexteP = tk.Tk()
    TexteP.wm_title("Distorsion de corriente (%)")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1A = tk.Label(TexteP, text = "En nuestro caso, Icc/IL = {:.1f}".format(R) + " y la tencion nominal vale {:.0f}".format(TN) + " V." +"\n" + " Por las harmonicas pares: " + "\n"\
                                   + "CAIMC(h<11) = {:.1f}".format(CAIMCP11) + "%" + "\n"+" CAIMC(11 ≤ h < 17) = {:.2f}".format(CAIMCP1117) + "\n"\
                                   + + "\n"+" CAIMC(17 ≤ h < 23) = {:.2f}".format(CAIMCP1723) + "%"\
                                   + "\n" + " CAIMC(23 ≤ h < 35) = {:.2f}".format(CAIMCP2335) + "%"\
                                   + "\n" + " CAIMC(h ≥ 35) = {:.2f}".format(CAIMCP35) + " %." + "\n" + " Por las harmonicas impares: " + "\n"\
                                   + "CAIMC(h<11) = {:.2f}".format(CAIMCI11) + " %" + "\n"\
                                   + "CAIMC(17 ≤ h < 23) = {:.2f}".format(CAIMCI1723)+ " %" + "\n"\
                                   + "CAIMC(23 ≤ h < 35) = {:.2f}".format(CAIMCI2335)+ " %" + "\n"\
                                   + "CAIMC(h ≥ 35) = {:.2f}".format(CAIMCI35) + " %" + "\n"\
                                   + " y DATD = {:.2f}".format(DATD) + " %.",foreground = "white", background = "#2B00FA")
    label1A['font'] = f_label
    label1A.pack() 
           
    if TN < 69000:
        if R < 20:
            if CAIMCP11 > 1:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label2A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label2A['font'] = f_label
                label2A.pack()        
            if CAIMCI11 > 4:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label3A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label3A['font'] = f_label
                label3A.pack()        
            if CAIMCP1117 > 0.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label4A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label4A['font'] = f_label
                label4A.pack()        
            if CAIMCI1117 > 2:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label5A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label5A['font'] = f_label
                label5A.pack()        
            if CAIMCP1723 > 1.5*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label6A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label6A['font'] = f_label
                label6A.pack()        
            if CAIMCI1723 > 1.5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label7A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label7A['font'] = f_label
                label7A.pack()        
            if CAIMCP2335 > 0.25*0.6:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label8A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label8A['font'] = f_label
                label8A.pack()        
            if CAIMCI2335 > 0.6:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label9A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label9A['font'] = f_label
                label9A.pack()        
            if CAIMCP35 > 0.3*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label10A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label10A['font'] = f_label
                label10A.pack()        
            if CAIMCI35 > 0.3:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label11A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label11A['font'] = f_label
                label11A.pack()            
            if DATD > 5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label12A = tk.Label(TexteP, text = "el DATD no respeta la especificacion.",foreground = "white", background = "#2B00FA")
                label12A['font'] = f_label
                label12A.pack()                            
            if CAIMCP11 <= 1 and CAIMCI11 <= 4 and CAIMCP1117 <= 0.5 and CAIMCI1117 <= 2 and CAIMCP1723 <= 1.5*0.25 and CAIMCI1723 <= 1.5 and CAIMCP2335 <= 0.25*0.6 and CAIMCI2335 <= 0.6 and CAIMCP35 <= 0.25*0.3 and CAIMCI35 <= 0.3 and DATD <= 5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label13A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label13A['font'] = f_label
                label13A.pack()                            
        if 20 <= R < 50:
            if CAIMCP11 > 0.25*7:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label14A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label14A['font'] = f_label
                label14A.pack()                            
            if CAIMCI11 > 7:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label15A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label15A['font'] = f_label
                label15A.pack()                            
            if CAIMCP1117 > 0.25*3.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label16A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label16A['font'] = f_label
                label16A.pack()                            
            if CAIMCI1117 > 3.5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label17A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label17A['font'] = f_label
                label17A.pack()                            
            if CAIMCP1723 > 2.5*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label18A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label18A['font'] = f_label
                label18A.pack()                            
            if CAIMCI1723 > 2.5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label19A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label19A['font'] = f_label
                label19A.pack()                            
            if CAIMCP2335 > 0.25*1:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label20A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label20A['font'] = f_label
                label20A.pack()                            
            if CAIMCI2335 > 1:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label21A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label21A['font'] = f_label
                label21A.pack()                            
            if CAIMCP35 > 0.5*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label22A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label22A['font'] = f_label
                label22A.pack()                            
            if CAIMCI35 > 0.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label23A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label23A['font'] = f_label
                label23A.pack()                            
            if DATD > 8:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label24A = tk.Label(TexteP, text = "el DATD no respeta la especificacion",foreground = "white", background = "#2B00FA")
                label24A['font'] = f_label
                label24A.pack()                            
            if CAIMCP11 <= 0.25*7 and CAIMCI11 <= 7 and CAIMCP1117 <= 0.25*3.5 and CAIMCI1117 <= 3.5 and CAIMCP1723 <= 2.5*0.25 and CAIMCI1723 <= 2.5 and CAIMCP2335 <= 0.25*1 and CAIMCI2335 <= 1 and CAIMCP35 <= 0.25*0.5 and CAIMCI35 <= 0.5 and DATD <= 8:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label25A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label25A['font'] = f_label
                label25A.pack()                                    
        if 50 <= R < 100:
            if CAIMCP11 > 0.25*10:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label26A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label26A['font'] = f_label
                label26A.pack()                                    
            if CAIMCI11 > 10:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label27A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label27A['font'] = f_label
                label27A.pack()                                    
            if CAIMCP1117 > 0.25*4.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label28A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label28A['font'] = f_label
                label28A.pack()                                    
            if CAIMCI1117 > 4.5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label29A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label29A['font'] = f_label
                label29A.pack()                                    
            if CAIMCP1723 > 4*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label30A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label30A['font'] = f_label
                label30A.pack()                                    
            if CAIMCI1723 > 4:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label31A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label31A['font'] = f_label
                label31A.pack()                                    
            if CAIMCP2335 > 0.25*1.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label32A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label32A['font'] = f_label
                label32A.pack()                                    
            if CAIMCI2335 > 1.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label33A = tk.Label(TexteP, text = "Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label33A['font'] = f_label
                label33A.pack()                                                    
            if CAIMCP35 > 0.7*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label34A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label34A['font'] = f_label
                label34A.pack()                                                    
            if CAIMCI35 > 0.7:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label35A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label35A['font'] = f_label
                label35A.pack()                                                    
            if DATD > 12:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label36A = tk.Label(TexteP, text = "el DATD no respeta la especificacion.",foreground = "white", background = "#2B00FA")
                label36A['font'] = f_label
                label36A.pack()                                                    
            if CAIMCP11 <= 0.25*10 and CAIMCI11 <= 10 and CAIMCP1117 <= 0.25*4.5 and CAIMCI1117 <= 4.5 and CAIMCP1723 <= 4*0.25 and CAIMCI1723 <= 4 and CAIMCP2335 <= 0.25*1.5 and CAIMCI2335 <= 1.5 and CAIMCP35 <= 0.25*0.7 and CAIMCI35 <= 0.7 and DATD <= 12:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label37A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label37A['font'] = f_label
                label37A.pack()                                                    
        if 100 <= R < 1000:
            if CAIMCP11 > 0.25*12:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label38A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label38A['font'] = f_label
                label38A.pack()                                                    
            if CAIMCI11 > 12:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label39A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label39A['font'] = f_label
                label39A.pack()                                                    
            if CAIMCP1117 > 0.25*5.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label40A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label40A['font'] = f_label
                label40A.pack()                                                    
            if CAIMCI1117 > 5.5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label41A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label41A['font'] = f_label
                label41A.pack()                                                    
            if CAIMCP1723 > 5*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label42A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label42A['font'] = f_label
                label42A.pack()                                                    
            if CAIMCI1723 > 5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label43A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label43A['font'] = f_label
                label43A.pack()                                                    
            if CAIMCP2335 > 0.25*2:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label44A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label44A['font'] = f_label
                label44A.pack()                                                    
            if CAIMCI2335 > 2:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label45A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label45A['font'] = f_label
                label45A.pack()                                                    
            if CAIMCP35 > 1*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label46A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label46A['font'] = f_label
                label46A.pack()                                                    
            if CAIMCI35 > 1:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label47A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label47A['font'] = f_label
                label47A.pack()                                                    
            if DATD > 15:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label48A = tk.Label(TexteP, text = "el DATD no respeta la especificacion",foreground = "white", background = "#2B00FA")
                label48A['font'] = f_label
                label48A.pack()                                                    
            if CAIMCP11 <= 0.25*12 and CAIMCI11 <= 12 and CAIMCP1117 <= 0.25*5.5 and CAIMCI1117 <= 5.5 and CAIMCP1723 <= 5*0.25 and CAIMCI1723 <= 5 and CAIMCP2335 <= 0.25*2 and CAIMCI2335 <= 2 and CAIMCP35 <= 0.25*1 and CAIMCI35 <= 1 and DATD <= 15:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label49A = tk.Label(TexteP, text = "Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label49A['font'] = f_label
                label49A.pack()                                                    
        if R >= 1000:
            if CAIMCP11 > 0.25*15:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label50A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label50A['font'] = f_label
                label50A.pack()
            if CAIMCI11 > 15:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label51A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label51A['font'] = f_label
                label51A.pack()
            if CAIMCP1117 > 0.25*7:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label52A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label52A['font'] = f_label
                label52A.pack()
            if CAIMCI1117 > 7:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label53A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label53A['font'] = f_label
                label53A.pack()                
            if CAIMCP1723 > 6*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label54A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label54A['font'] = f_label
                label54A.pack()                
            if CAIMCI1723 > 6:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label55A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label55A['font'] = f_label
                label55A.pack()                
            if CAIMCP2335 > 0.25*2.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label56A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label56A['font'] = f_label
                label56A.pack()                
            if CAIMCI2335 > 2.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label57A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label57A['font'] = f_label
                label57A.pack()                
            if CAIMCP35 > 1.4*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label58A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label58A['font'] = f_label
                label58A.pack()                
            if CAIMCI35 > 1.4:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label59A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label59A['font'] = f_label
                label59A.pack()                
            if DATD > 20:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label60A = tk.Label(TexteP, text = "el DATD no respeta la especificacion.",foreground = "white", background = "#2B00FA")
                label60A['font'] = f_label
                label60A.pack()                
            if CAIMCP11 <= 0.25*15 and CAIMCI11 <= 15 and CAIMCP1117 <= 0.25*7 and CAIMCI1117 <= 7 and CAIMCP1723 <= 6*0.25 and CAIMCI1723 <= 6 and CAIMCP2335 <= 0.25*2.5 and CAIMCI2335 <= 2.5 and CAIMCP35 <= 0.25*1.4 and CAIMCI35 <= 1.4 and DATD <= 20:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label61A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label61A['font'] = f_label
                label61A.pack()                   
    if 69000 <= TN < 161000:
        if R > 20:
            if CAIMCP11 > 0.25*2:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label62A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label62A['font'] = f_label
                label62A.pack()
            if CAIMCI11 > 2:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label63A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label63A['font'] = f_label
                label63A.pack()
            if CAIMCP1117 > 0.25*1:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label64A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label64A['font'] = f_label
                label64A.pack()
            if CAIMCI1117 > 1:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label65A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label65A['font'] = f_label
                label65A.pack()
            if CAIMCP1723 > 0.75*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label66A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label66A['font'] = f_label
                label66A.pack()
            if CAIMCI1723 > 0.75:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label67A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label67A['font'] = f_label
                label67A.pack()
            if CAIMCP2335 > 0.25*0.3:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label68A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label68A['font'] = f_label
                label68A.pack()
            if CAIMCI2335 > 0.3:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label69A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label69A['font'] = f_label
                label69A.pack()
            if CAIMCP35 > 0.15*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label70A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label70A['font'] = f_label
                label70A.pack()
            if CAIMCI35 > 0.15:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label71A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label71A['font'] = f_label
                label71A.pack()
            if DATD > 2.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label72A = tk.Label(TexteP, text = "el DATD no respeta la especificacion.",foreground = "white", background = "#2B00FA")
                label72A['font'] = f_label
                label72A.pack()
            if CAIMCP11 <= 0.25*2 and CAIMCI11 <= 2 and CAIMCP1117 <= 0.25*1 and CAIMCI1117 <= 1 and CAIMCP1723 <= 0.75*0.25 and CAIMCI1723 <= 0.75 and CAIMCP2335 <= 0.25*0.3 and CAIMCI2335 <= 0.3 and CAIMCP35 <= 0.25*0.15 and CAIMCI35 <= 0.15 and DATD <= 2.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label73A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label73A['font'] = f_label
                label73A.pack()
        if 20 <= R < 50:
            if CAIMCP11 > 0.25*3.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label74A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label74A['font'] = f_label
                label74A.pack()
            if CAIMCI11 > 3.5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label75A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label75A['font'] = f_label
                label75A.pack()
            if CAIMCP1117 > 0.25*1.75:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label76A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label76A['font'] = f_label
                label76A.pack()
            if CAIMCI1117 > 1.75:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label77A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label77A['font'] = f_label
                label77A.pack()
            if CAIMCP1723 > 1.25*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label78A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label78A['font'] = f_label
                label78A.pack()
            if CAIMCI1723 > 1.25:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label79A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label79A['font'] = f_label
                label79A.pack()
            if CAIMCP2335 > 0.25*0.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label80A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label80A['font'] = f_label
                label80A.pack()
            if CAIMCI2335 > 0.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label81A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label81A['font'] = f_label
                label81A.pack()
            if CAIMCP35 > 0.25*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label82A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label82A['font'] = f_label
                label82A.pack()
            if CAIMCI35 > 0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label83A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label83A['font'] = f_label
                label83A.pack()
            if DATD > 4:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label84A = tk.Label(TexteP, text = "el DATD no respeta la especificacion.",foreground = "white", background = "#2B00FA")
                label84A['font'] = f_label
                label84A.pack()
            if CAIMCP11 <= 0.25*3.5 and CAIMCI11 <= 3.5 and CAIMCP1117 <= 0.25*1.75 and CAIMCI1117 <= 1.75 and CAIMCP1723 <= 1.25*0.25 and CAIMCI1723 <= 1.25 and CAIMCP2335 <= 0.25*0.5 and CAIMCI2335 <= 0.5 and CAIMCP35 <= 0.25*0.25 and CAIMCI35 <= 0.25 and DATD <= 4:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label85A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label85A['font'] = f_label
                label85A.pack()
        if 50 <= R < 100:
            if CAIMCP11 > 0.25*5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label86A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label86A['font'] = f_label
                label86A.pack()
            if CAIMCI11 > 5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label87A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label87A['font'] = f_label
                label87A.pack()
            if CAIMCP1117 > 0.25*2.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label88A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label88A['font'] = f_label
                label88A.pack()
            if CAIMCI1117 > 2.25:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label89A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label89A['font'] = f_label
                label89A.pack()
            if CAIMCP1723 > 2*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label90A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label90A['font'] = f_label
                label90A.pack()
            if CAIMCI1723 > 2:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label91A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label91A['font'] = f_label
                label91A.pack()
            if CAIMCP2335 > 0.25*0.75:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label92A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label92A['font'] = f_label
                label92A.pack()
            if CAIMCI2335 > 0.75:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label93A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label93A['font'] = f_label
                label93A.pack()
            if CAIMCP35 > 0.25*0.35:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label94A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label94A['font'] = f_label
                label94A.pack()
            if CAIMCI35 > 0.35:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label95A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label95A['font'] = f_label
                label95A.pack()
            if DATD > 6:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label96A = tk.Label(TexteP, text = "el DATD no respeta la especificacion.",foreground = "white", background = "#2B00FA")
                label96A['font'] = f_label
                label96A.pack()
            if CAIMCP11 <= 0.25*5 and CAIMCI11 <= 5 and CAIMCP1117 <= 0.25*2.25 and CAIMCI1117 <= 2.25 and CAIMCP1723 <= 2*0.25 and CAIMCI1723 <= 2 and CAIMCP2335 <= 0.25*0.75 and CAIMCI2335 <= 0.75 and CAIMCP35 <= 0.25*0.35 and CAIMCI35 <= 0.35 and DATD <= 6:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label97A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label97A['font'] = f_label
                label97A.pack()
        if 100 <= R < 1000:
            if CAIMCP11 > 0.25*6:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label98A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label98A['font'] = f_label
                label98A.pack()
            if CAIMCI11 > 6:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label99A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label99A['font'] = f_label
                label99A.pack()
            if CAIMCP1117 > 0.25*2.75:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label100A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label100A['font'] = f_label
                label100A.pack()
            if CAIMCI1117 > 2.75:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label101A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label101A['font'] = f_label
                label101A.pack()
            if CAIMCP1723 > 2.5*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label102A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label102A['font'] = f_label
                label102A.pack()
            if CAIMCI1723 > 2.5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label103A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label103A['font'] = f_label
                label103A.pack()
            if CAIMCP2335 > 0.25*1:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label104A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label104A['font'] = f_label
                label104A.pack()
            if CAIMCI2335 > 1:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label105A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label105A['font'] = f_label
                label105A.pack()
            if CAIMCP35 > 0.25*0.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label106A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label106A['font'] = f_label
                label106A.pack()
            if CAIMCI35 > 0.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label107A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label107A['font'] = f_label
                label107A.pack()
            if DATD > 7.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label108A = tk.Label(TexteP, text = "el DATD no respeta la especificacion.",foreground = "white", background = "#2B00FA")
                label108A['font'] = f_label
                label108A.pack()
            if CAIMCP11 <= 0.25*6 and CAIMCI11 <= 6 and CAIMCP1117 <= 0.25*2.75 and CAIMCI1117 <= 2.75 and CAIMCP1723 <= 0.25*2.5 and CAIMCI1723 <= 2.5 and CAIMCP2335 <= 0.25*1 and CAIMCI2335 <= 1 and CAIMCP35 <= 0.25*0.5 and CAIMCI35 <= 0.5 and DATD <= 7.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label109A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label109A['font'] = f_label
                label109A.pack()
        if R >= 1000:
            if CAIMCP11 > 0.25*7.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label110A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label110A['font'] = f_label
                label110A.pack()
            if CAIMCI11 > 7.5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label111A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label111A['font'] = f_label
                label111A.pack()
            if CAIMCP1117 > 0.25*3.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label112A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label112A['font'] = f_label
                label112A.pack()
            if CAIMCI1117 > 3.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label113A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label113A['font'] = f_label
                label113A.pack()
            if CAIMCP1723 > 3*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label114A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label114A['font'] = f_label
                label114A.pack()                
            if CAIMCI1723 > 3:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label115A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label115A['font'] = f_label
                label115A.pack() 
            if CAIMCP2335 > 0.25*1.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label116A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label116A['font'] = f_label
                label116A.pack() 
            if CAIMCI2335 > 1.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label117A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label117A['font'] = f_label
                label117A.pack() 
            if CAIMCP35 > 0.25*0.7:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label118A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label118A['font'] = f_label
                label118A.pack() 
            if CAIMCI35 > 0.7:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label119A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label119A['font'] = f_label
                label119A.pack() 
            if DATD > 10:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label120A = tk.Label(TexteP, text = "el DATD no respeta la especificacion.",foreground = "white", background = "#2B00FA")
                label120A['font'] = f_label
                label120A.pack() 
            if CAIMCP11 <= 0.25*7.5 and CAIMCI11 <= 7.5 and CAIMCP1117 <= 0.25*3.5 and CAIMCI1117 <= 3.5 and CAIMCP1723 <= 3*0.25 and CAIMCI1723 <= 3 and CAIMCP2335 <= 0.25*1.25 and CAIMCI2335 <= 1.25 and CAIMCP35 <= 0.25*0.7 and CAIMCI35 <= 0.7 and DATD <= 10:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label121A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label121A['font'] = f_label
                label121A.pack()
    if TN >= 161000:
        if R < 50:
            if CAIMCP11 > 0.25*2:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label122A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label122A['font'] = f_label
                label122A.pack()
            if CAIMCI11 > 2:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label123A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label123A['font'] = f_label
                label123A.pack()
            if CAIMCP1117 > 0.25*1:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label124A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label124A['font'] = f_label
                label124A.pack()
            if CAIMCI1117 > 1:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label125A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label125A['font'] = f_label
                label125A.pack()
            if CAIMCP1723 > 0.75*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label126A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label126A['font'] = f_label
                label126A.pack()
            if CAIMCI1723 > 0.75:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label127A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label127A['font'] = f_label
                label127A.pack()
            if CAIMCP2335 > 0.25*0.3:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label128A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label128A['font'] = f_label
                label128A.pack()
            if CAIMCI2335 > 0.3:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label129A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label129A['font'] = f_label
                label129A.pack()
            if CAIMCP35 > 0.15*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label130A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label130A['font'] = f_label
                label130A.pack()
            if CAIMCI35 > 0.15:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label131A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label131A['font'] = f_label
                label131A.pack()
            if DATD > 2.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label132A = tk.Label(TexteP, text = "el DATD no respeta la especificacion.",foreground = "white", background = "#2B00FA")
                label132A['font'] = f_label
                label132A.pack()
            if CAIMCP11 <= 0.25*2 and CAIMCI11 <= 2 and CAIMCP1117 <= 0.25*1 and CAIMCI1117 <= 1 and CAIMCP1723 <= 0.75*0.25 and CAIMCI1723 <= 0.75 and CAIMCP2335 <= 0.25*0.3 and CAIMCI2335 <= 0.3 and CAIMCP35 <= 0.25*0.15 and CAIMCI35 <= 0.15 and DATD <= 2.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label133A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label133A['font'] = f_label
                label133A.pack()
        if R >= 50:
            if CAIMCP11 > 0.25*3:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label134A = tk.Label(TexteP, text = "las armonicas pares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label134A['font'] = f_label
                label134A.pack()
            if CAIMCI11 > 3:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label135A = tk.Label(TexteP, text = "las armonicas impares de rango inferior a 11 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label135A['font'] = f_label
                label135A.pack()
            if CAIMCP1117 > 0.25*1.5:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label136A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label136A['font'] = f_label
                label136A.pack()
            if CAIMCI1117 > 1.5:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label137A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label137A['font'] = f_label
                label137A.pack()
            if CAIMCP1723 > 1.15*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label138A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label138A['font'] = f_label
                label138A.pack()
            if CAIMCI1723 > 1.15:    
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label139A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label139A['font'] = f_label
                label139A.pack()
            if CAIMCP2335 > 0.25*0.45:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label140A = tk.Label(TexteP, text = "las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label140A['font'] = f_label
                label140A.pack()
            if CAIMCI2335 > 0.45:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label141A = tk.Label(TexteP, text = "las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.",foreground = "white", background = "#2B00FA")
                label141A['font'] = f_label
                label141A.pack()
            if CAIMCP35 > 0.22*0.25:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label142A = tk.Label(TexteP, text = "las armonicas pares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label142A['font'] = f_label
                label142A.pack()
            if CAIMCI35 > 0.22:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label143A = tk.Label(TexteP, text = "las armonicas impares de rango superior a 35 no respetan la especification.",foreground = "white", background = "#2B00FA")
                label143A['font'] = f_label
                label143A.pack()
            if DATD > 3.75:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label144A = tk.Label(TexteP, text = "el DATD no respeta la especificacion.",foreground = "white", background = "#2B00FA")
                label144A['font'] = f_label
                label144A.pack()
            if CAIMCP11 <= 0.25*3 and CAIMCI11 <= 3 and CAIMCP1117 <= 0.25*1.5 and CAIMCI1117 <= 1.5 and CAIMCP1723 <= 1.15*0.25 and CAIMCI1723 <= 1.15 and CAIMCP2335 <= 0.25*0.45 and CAIMCI2335 <= 0.45 and CAIMCP35 <= 0.25*0.22 and CAIMCI35 <= 0.22 and DATD <= 3.75:
                f_label = font.Font(family='Times New Roman', size=10)
                f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
                label145A = tk.Label(TexteP, text = "el DATD y el CAIMC de todas las armonicas cumplen la especification.",foreground = "white", background = "#2B00FA")
                label145A['font'] = f_label
                label145A.pack()

def DistI(fichier):
    L = extractiondonne(fichier)
    NB = []
    H = []
    DATD = 0
    A = []
    R = Icc/Inom
    Caim = []
    hihi = []
    for i in range(1,len(L)):
        NB.append(i)
    for i in range (106+3*49+3,106+3*49+3 + 3*49 + 1,3):
        for j in range(1,len(L)):
            A.append(L[j][i])
        H.append(mediane(A))
        A = []
    HP11 = [H[0],H[2],H[4],H[6],H[8],H[10]]
    HI11 = [H[3],H[5],H[7],H[9]]
    HP1117 = [H[12],H[14],H[16]]
    HI1117 = [H[11],H[13],H[15]]
    HP1723 = [H[18],H[20],H[22]]
    HI1723 = [H[17],H[19],H[21]]
    HP2335 = [H[24],H[26],H[28],H[30],H[32],H[34]]
    HI2335 = [H[23],H[25],H[27],H[29],H[31],H[33]]
    HP35 = [H[36],H[38],H[40],H[42],H[44],H[46],H[48]]
    HI35 = [H[35],H[37],H[39],H[41],H[43],H[45],H[47],H[49]]    
    for i in range(2,50):
        DATD += (H[i]**2)/(H[1]**2)
    DATD = 100*sqrt(DATD)
    for i in range(len(H)):
        Caim.append(100*H[i]/H[1])
    CAIMCP11 = 100*(max(HP11)/H[1])
    CAIMCI11 = 100*(max(HI11)/H[1])
    CAIMCP1117 = 100*(max(HP1117)/H[1])
    CAIMCI1117 = 100*(max(HI1117)/H[1])
    CAIMCP1723 = 100*(max(HP1723)/H[1])
    CAIMCI1723 = 100*(max(HI1723)/H[1])
    CAIMCP2335 = 100*(max(HP2335)/H[1])
    CAIMCI2335 = 100*(max(HI2335)/H[1])
    CAIMCP35 = 100*(max(HP35)/H[1])
    CAIMCI35 = 100*(max(HI35)/H[1])
    document.add_heading("Distorsión de corriente: ",1)
    document.add_paragraph("Recordamos que la especificacion CFE-L000045 define dos valores fondamentales"\
                           + "para evaluar la distorción de corriente: el CAIMC y el DATD. Esos valores"\
                           + " son definidos asi:" + "\n")
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\CAIMCDATD.png", width=Inches(4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #document.add_paragraph("En nuestro caso, CAIMT = {:.2f}".format(CAIMT) + " y DATT = {:.2f}".format(DATT))
    if TN < 69000:
        document.add_paragraph("La especificacion CFE-L000045 especifica las siguientes tolerancias sobre esos valores " + "\n")    
        document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\Tolérance distortion en courant moins 69 kV.png", width=Inches(4))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph("En nuestro caso, Icc/IL = {:.1f}".format(R) + " y la tencion nominal vale {:.0f}".format(TN) + " V." +"\n" + " Por las harmonicas pares: " + "\n"\
                                   + "- CAIMC(h<11) = {:.1f}".format(CAIMCP11) + "%" + "\n" + "CAIMC(11 ≤ h < 17) = {:.2f}".format(CAIMCP1117) +"%" +"\n"\
                                   + "- CAIMC(17 ≤ h < 23) = {:.2f}".format(CAIMCP1723) + "%" +"\n"\
                                   + "- CAIMC(23 ≤ h < 35) = {:.2f}".format(CAIMCP2335) + "%" +"\n"\
                                   + "- CAIMC(h ≥ 35) = {:.2f}".format(CAIMCP35) + " %" + "\n" + " Por las harmonicas impares: " + "\n"\
                                   + "- CAIMC(h<11) = {:.2f}".format(CAIMCI11) + " %" + "\n"\
                                   + "- CAIMC(17 ≤ h < 23) = {:.2f}".format(CAIMCI1723)+ " %" + "\n"\
                                   + "- CAIMC(23 ≤ h < 35) = {:.2f}".format(CAIMCI2335)+ " %" + "\n"\
                                   + "- CAIMC(h ≥ 35) = {:.2f}".format(CAIMCI35) + " %" + "\n"\
                                   + " Y, además:  DATD = {:.2f}".format(DATD) + " %." + "\n")
        if R < 20:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 4:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 4:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 2:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 2:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 1.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 0.6:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.6:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 0.3:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.3:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                    
            if CAIMCP11 > 1:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 4:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.5:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 2:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 1.5*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 1.5:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*0.6:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 0.6:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 0.3*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 0.3:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 5:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion.")
            if CAIMCP11 <= 1 and CAIMCI11 <= 4 and CAIMCP1117 <= 0.5 and CAIMCI1117 <= 2 and CAIMCP1723 <= 1.5*0.25 and CAIMCI1723 <= 1.5 and CAIMCP2335 <= 0.25*0.6 and CAIMCI2335 <= 0.6 and CAIMCP35 <= 0.25*0.3 and CAIMCI35 <= 0.3 and DATD <= 5:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")
        if 20 <= R < 50:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 7:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 7:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 3.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 3.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 2.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 2.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 1:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 0.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            
            if CAIMCP11 > 0.25*7:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 7:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*3.5:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 3.5:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 2.5*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 2.5:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*1:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 1:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 0.5*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 0.5:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 8:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion")
            if CAIMCP11 <= 0.25*7 and CAIMCI11 <= 7 and CAIMCP1117 <= 0.25*3.5 and CAIMCI1117 <= 3.5 and CAIMCP1723 <= 2.5*0.25 and CAIMCI1723 <= 2.5 and CAIMCP2335 <= 0.25*1 and CAIMCI2335 <= 1 and CAIMCP35 <= 0.25*0.5 and CAIMCI35 <= 0.5 and DATD <= 8:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")
        
        if 50 <= R < 100:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 10:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 10:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 4.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 4.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 4:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 4:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 1.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 0.7:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.7:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            
            if CAIMCP11 > 0.25*10:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 10:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*4.5:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 4.5:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 4*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 4:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*1.5:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 1.5:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 0.7*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 0.7:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 12:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion.")
            if CAIMCP11 <= 0.25*10 and CAIMCI11 <= 10 and CAIMCP1117 <= 0.25*4.5 and CAIMCI1117 <= 4.5 and CAIMCP1723 <= 4*0.25 and CAIMCI1723 <= 4 and CAIMCP2335 <= 0.25*1.5 and CAIMCI2335 <= 1.5 and CAIMCP35 <= 0.25*0.7 and CAIMCI35 <= 0.7 and DATD <= 12:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")

        if 100 <= R < 1000:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 12:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 12:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 5.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 5.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 2:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 2:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 1:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            
            if CAIMCP11 > 0.25*12:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 12:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*5.5:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 5.5:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 5*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 5:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*2:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 2:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 1*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 1:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 15:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion")
            if CAIMCP11 <= 0.25*12 and CAIMCI11 <= 12 and CAIMCP1117 <= 0.25*5.5 and CAIMCI1117 <= 5.5 and CAIMCP1723 <= 5*0.25 and CAIMCI1723 <= 5 and CAIMCP2335 <= 0.25*2 and CAIMCI2335 <= 2 and CAIMCP35 <= 0.25*1 and CAIMCI35 <= 1 and DATD <= 15:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")
        if R >= 1000:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 15:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 15:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 7:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 7:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 6:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 6:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 2.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 2.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 1.4:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1.4:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            
            if CAIMCP11 > 0.25*15:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 15:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*7:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 7:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 6*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 6:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*2.5:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 2.5:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 1.4*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 1.4:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 20:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion.")
            if CAIMCP11 <= 0.25*15 and CAIMCI11 <= 15 and CAIMCP1117 <= 0.25*7 and CAIMCI1117 <= 7 and CAIMCP1723 <= 6*0.25 and CAIMCI1723 <= 6 and CAIMCP2335 <= 0.25*2.5 and CAIMCI2335 <= 2.5 and CAIMCP35 <= 0.25*1.4 and CAIMCI35 <= 1.4 and DATD <= 20:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")    
    if 69000 <= TN < 161000:
        document.add_paragraph("La especificacion CFE-L000045 especifica las siguientes tolerancias sobre esos valores: ")    
        document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\Tolerance distortion courant 69kV a 161kV.png", width=Inches(4))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph("En nuestro caso, Icc/IL = {:.1f}".format(R) + " y la tencion nominal vale {:.0f}".format(TN) + " V." + "\n" + " Por las harmonicas pares: " + "\n"\
                                   + "CAIMC(h<11) = {:.2f}".format(CAIMCP11) + "\n" + " CAIMC(11 ≤ h < 17) = {:.2f}".format(CAIMCP1117) + "%" + "\n"\
                                   + " CAIMC(17 ≤ h < 23) = {:.2f}".format(CAIMCP1723) +"%" +"\n"\
                                   + " %, CAIMC(23 ≤ h < 35) = {:.2f}".format(CAIMCP2335) +"%" +"\n"\
                                   + " %, CAIMC(h ≥ 35) = {:.2f}".format(CAIMCP35) + " %." + "\n" + " Por las harmonicas impares: " + "\n"\
                                   + "CAIMC(h<11) = {:.2f}".format(CAIMCI11) + " %" + "\n"\
                                   + "CAIMC(17 ≤ h < 23) = {:.2f}".format(CAIMCI1723)+ " %" + "\n"\
                                   + "CAIMC(23 ≤ h < 35) = {:.2f}".format(CAIMCI2335)+ " %" + "\n"\
                                   + "CAIMC(h ≥ 35) = {:.2f}".format(CAIMCI35) + " %" + "\n"\
                                   + " y DATD = {:.2f}".format(DATD) + " %." + "\n")
        if R > 20:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 2:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 2:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 1:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 0.75:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.75:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 0.3:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.3:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 0.15:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.15:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            
            if CAIMCP11 > 0.25*2:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 2:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*1:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 1:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 0.75*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 0.75:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*0.3:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 0.3:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 0.15*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 0.15:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 2.5:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion.")
            if CAIMCP11 <= 0.25*2 and CAIMCI11 <= 2 and CAIMCP1117 <= 0.25*1 and CAIMCI1117 <= 1 and CAIMCP1723 <= 0.75*0.25 and CAIMCI1723 <= 0.75 and CAIMCP2335 <= 0.25*0.3 and CAIMCI2335 <= 0.3 and CAIMCP35 <= 0.25*0.15 and CAIMCI35 <= 0.15 and DATD <= 2.5:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")
        if 20 <= R < 50:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 3.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 3.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 1.75:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1.75:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 1.25:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1.25:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 0.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 0.25:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.25:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            
            if CAIMCP11 > 0.25*3.5:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 3.5:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*1.75:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 1.75:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 1.25*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 1.25:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*0.5:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 0.5:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 0.25*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 0.25:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 4:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion.")
            if CAIMCP11 <= 0.25*3.5 and CAIMCI11 <= 3.5 and CAIMCP1117 <= 0.25*1.75 and CAIMCI1117 <= 1.75 and CAIMCP1723 <= 1.25*0.25 and CAIMCI1723 <= 1.25 and CAIMCP2335 <= 0.25*0.5 and CAIMCI2335 <= 0.5 and CAIMCP35 <= 0.25*0.25 and CAIMCI35 <= 0.25 and DATD <= 4:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")
        if 50 <= R < 100:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 2.25:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 2.25:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 2:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 2:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 0.75:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.75:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 0.35:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.35:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            
            if CAIMCP11 > 0.25*5:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 5:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*2.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 2.25:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 2*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 2:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*0.75:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 0.75:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 0.25*0.35:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 0.35:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 6:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion.")
            if CAIMCP11 <= 0.25*5 and CAIMCI11 <= 5 and CAIMCP1117 <= 0.25*2.25 and CAIMCI1117 <= 2.25 and CAIMCP1723 <= 2*0.25 and CAIMCI1723 <= 2 and CAIMCP2335 <= 0.25*0.75 and CAIMCI2335 <= 0.75 and CAIMCP35 <= 0.25*0.35 and CAIMCI35 <= 0.35 and DATD <= 6:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")
        if 100 <= R < 1000:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 6:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 6:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 2.75:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 2.75:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 2.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 2.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 1:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 0.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            
            if CAIMCP11 > 0.25*6:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 6:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*2.75:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 2.75:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 2.5*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 2.5:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*1:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 1:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 0.25*0.5:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 0.5:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 7.5:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion.")
            if CAIMCP11 <= 0.25*6 and CAIMCI11 <= 6 and CAIMCP1117 <= 0.25*2.75 and CAIMCI1117 <= 2.75 and CAIMCP1723 <= 0.25*2.5 and CAIMCI1723 <= 2.5 and CAIMCP2335 <= 0.25*1 and CAIMCI2335 <= 1 and CAIMCP35 <= 0.25*0.5 and CAIMCI35 <= 0.5 and DATD <= 7.5:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")
        if R >= 1000:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 7.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 7.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 3.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 3.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 3:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 3:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 1.25:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1.25:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 0.7:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.7:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            if CAIMCP11 > 0.25*7.5:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 7.5:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*3.5:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 3.5:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 3*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 3:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*1.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 1.25:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 0.25*0.7:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 0.7:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 10:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion.")
            if CAIMCP11 <= 0.25*7.5 and CAIMCI11 <= 7.5 and CAIMCP1117 <= 0.25*3.5 and CAIMCI1117 <= 3.5 and CAIMCP1723 <= 3*0.25 and CAIMCI1723 <= 3 and CAIMCP2335 <= 0.25*1.25 and CAIMCI2335 <= 1.25 and CAIMCP35 <= 0.25*0.7 and CAIMCI35 <= 0.7 and DATD <= 10:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")
    if TN >= 161000:
        document.add_paragraph("La especificacion CFE-L000045 especifica las siguientes tolerancias sobre esos valores: ")    
        document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\Tolerance distortion courant sup a 161 kV.png", width=Inches(4))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph("En nuestro caso, Icc/IL = {:.1f}".format(R) + " y la tencion nominal vale {:.0f}".format(TN) + " V. Por las harmonicas pares: " + "\n"\
                                   + "CAIMC(h<11) = {:.2f}".format(CAIMCP11) +"%"+ "\n" + " CAIMC(11 ≤ h < 17) = {:.2f}".format(CAIMCP1117) + "%" + "\n"\
                                   + " CAIMC(17 ≤ h < 23) = {:.2f}".format(CAIMCP1723) +"%"+ "\n" +\
                                   + " CAIMC(23 ≤ h < 35) = {:.2f}".format(CAIMCP2335) +"%"+ "\n" +\
                                   + " %, CAIMC(h ≥ 35) = {:.2f}".format(CAIMCP35) + " %." + "\n" + " Por las harmonicas impares: " + "\n"\
                                   + "CAIMC(h<11) = {:.2f}".format(CAIMCI11) + " %" + "\n"\
                                   + "CAIMC(17 ≤ h < 23) = {:.2f}".format(CAIMCI1723)+ " %" + "\n"\
                                   + "CAIMC(23 ≤ h < 35) = {:.2f}".format(CAIMCI2335)+ " %" + "\n"\
                                   + "CAIMC(h ≥ 35) = {:.2f}".format(CAIMCI35) + " %" + "\n"\
                                   + " y DATD = {:.2f}".format(DATD) + " %." + "\n")
        if R < 50:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 2:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 2:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 1:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 0.75:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.75:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 0.3:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.3:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 0.15:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.15:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            
            if CAIMCP11 > 0.25*2:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 2:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*1:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 1:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 0.75*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 0.75:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*0.3:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 0.3:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 0.15*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 0.15:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 2.5:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion.")
            if CAIMCP11 <= 0.25*2 and CAIMCI11 <= 2 and CAIMCP1117 <= 0.25*1 and CAIMCI1117 <= 1 and CAIMCP1723 <= 0.75*0.25 and CAIMCI1723 <= 0.75 and CAIMCP2335 <= 0.25*0.3 and CAIMCI2335 <= 0.3 and CAIMCP35 <= 0.25*0.15 and CAIMCI35 <= 0.15 and DATD <= 2.5:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")
        if R >= 50:
            for i in range(len(H)):
                if i < 11 and i != 1:
                    if Caim[i] <= 3:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 3:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i == 1:
                    hihi.append([i,"{:.1f}".format(Caim[i]),"X"])
                if 11 <= i < 17:
                    if Caim[i] <= 1.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1.5:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 17 <= i < 23:
                    if Caim[i] <= 1.15:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 1.15:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if 23 <= i < 35:
                    if Caim[i] <= 0.45:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.45:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
                if i >= 35:
                    if Caim[i] <= 0.22:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"Sí"])    
                    if Caim[i] > 0.22:
                        hihi.append([i,"{:.1f}".format(Caim[i]),"No"])                  
            
            if CAIMCP11 > 0.25*3:
                document.add_paragraph("Vemos que las armonicas pares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCI11 > 3:    
                document.add_paragraph("Vemos que las armonicas impares de rango inferior a 11 no respetan la especificacion.")
            if CAIMCP1117 > 0.25*1.5:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCI1117 > 1.5:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 11 y 17 no respetan la especificacion.")
            if CAIMCP1723 > 1.15*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCI1723 > 1.15:    
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 17 y 23 no respetan la especificacion.")
            if CAIMCP2335 > 0.25*0.45:
                document.add_paragraph("Vemos que las armonicas pares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCI2335 > 0.45:
                document.add_paragraph("Vemos que las armonicas impares de rango incluido entre 23 y 35 no respetan la especificacion.")
            if CAIMCP35 > 0.22*0.25:
                document.add_paragraph("Vemos que las armonicas pares de rango superior a 35 no respetan la especification.")
            if CAIMCI35 > 0.22:
                document.add_paragraph("Vemos que las armonicas impares de rango superior a 35 no respetan la especification.")
            if DATD > 3.75:
                document.add_paragraph("Vemos que el DATD no respeta la especificacion.")
            if CAIMCP11 <= 0.25*3 and CAIMCI11 <= 3 and CAIMCP1117 <= 0.25*1.5 and CAIMCI1117 <= 1.5 and CAIMCP1723 <= 1.15*0.25 and CAIMCI1723 <= 1.15 and CAIMCP2335 <= 0.25*0.45 and CAIMCI2335 <= 0.45 and CAIMCP35 <= 0.25*0.22 and CAIMCI35 <= 0.22 and DATD <= 3.75:
                document.add_paragraph("Como le podemos ver, el DATD y el CAIMC de todas las armonicas cumplen la especification.")
    document.add_paragraph("Para más detailles sobre las armonicas, tenemos los datos de la harmonicas en la siguiente tabla: " + "\n")
    menuTable = document.add_table(rows=1,cols=3)
    #menuTable.style = "Table grid"
    hdr_Cells = menuTable.rows[0].cells
    hdr_Cells[0].text = "número"
    hdr_Cells[1].text = "Ih/I1(%)"
    hdr_Cells[2].text = "¿cumple?"
    for a,b,c in hihi:
        row_Cells = menuTable.add_row().cells
        row_Cells[0].text = str(a)
        row_Cells[1].text = str(b)
        row_Cells[2].text = c
def BoutonFlicker(fichier):
    L = extractiondonne(fichier)
    NB = []
    APLT = []
    BPLT = []
    CPLT = []
    APST = []
    BPST = []
    CPST = []
    for i in range(1,len(L)):
        NB.append(i)
    for i in range (1,len(L)):
        APLT.append(L[i][-3])
        BPLT.append(L[i][-2])
        CPLT.append(L[i][-1])
        APST.append(L[i][-6])
        BPST.append(L[i][-5])
        CPST.append(L[i][-4])
    # Grafico del Plt
    FenetreP = tk.Tk()
    FenetreP.wm_title("Plt")
    FenetreP.configure(background="#2B00FA")
    fig = Figure(figsize=(6, 4), dpi=96)
    fag = Figure(figsize = (6,4), dpi = 95)
    ax = fig.add_subplot(111)
    ax.plot(NB,APLT, label = r"$Plt_A$")
    ax.legend()
    ax.plot(NB,BPLT, label = r"$Plt_B$")
    ax.legend()
    ax.plot(NB,CPLT, label = r"$Plt_C$")
    ax.legend()
    graph = FigureCanvasTkAgg(fig, master=FenetreP)
    canvas = graph.get_tk_widget()
    canvas.grid(row=0, column=0)
    #ax.show()
    # Grafico del Pst
    FenetreA = tk.Tk()
    FenetreA.wm_title("Pst")
    FenetreA.configure(background="#2B00FA")
    #fig = Figure(figsize=(6, 4), dpi=96)
    bx = fag.add_subplot(111)
    bx.plot(NB,APST, label = r"$Pst_A$")
    bx.legend()
    bx.plot(NB,BPST, label = r"$Pst_B$")
    bx.legend()
    bx.plot(NB,CPST, label = r"$Pst_C$")
    bx.legend()
    grapho = FigureCanvasTkAgg(fag, master=FenetreA)
    canvasa = grapho.get_tk_widget()
    canvasa.grid(row=0, column=0)
    #ax.show()
    # Zone de texte
    TexteP = tk.Tk()
    TexteP.wm_title(" Variaciones de tension")
    TexteP.configure(background="#2B00FA")
    #message
    f_label = font.Font(family='Times New Roman', size=10)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label1 = tk.Label(TexteP, text = "<Plt_A> = {:.2f}".format(moyenneA(APLT)) + "  "\
                      + "<Plt_B> =  {:.2f}".format(moyenneA(BPLT)) + "  "\
                      + "<Plt_C> = {:.2f}".format(moyenneA(CPLT))  + "  ",foreground = "white", background = "#2B00FA")
    label1['font'] = f_label
    label1.pack()
    label1A = tk.Label(TexteP, text = "<Pst_A> = {:.2f}".format(moyenneA(APST)) + "  "\
                      + "<Pst_B> =  {:.2f}".format(moyenneA(BPST)) + "  "\
                      + "<Pst_C> = {:.2f}".format(moyenneA(CPST))  + "  ",foreground = "white", background = "#2B00FA")
    label1A['font'] = f_label
    label1A.pack()    
    label2 = tk.Label(TexteP, text= "Plt_A_mediana = {:.2f}".format(mediane(APLT)) + "  " +\
                                    "Plt_B_mediana = {:.2f}".format(mediane(BPLT)) + "  " +\
                                    "Plt_C_mediana = {:.2f}".format(mediane(CPLT)),foreground = "white", background = "#2B00FA")
                                                                           
    label2['font'] = f_label
    label2.pack()
    label2A = tk.Label(TexteP, text= "Pst_A_mediana = {:.2f}".format(mediane(APST)) + "  " +\
                                    "Pst_B_mediana = {:.2f}".format(mediane(BPST)) + "  " +\
                                    "Pst_C_mediana = {:.2f}".format(mediane(CPST)),foreground = "white", background = "#2B00FA")
                                                                           
    label2A['font'] = f_label
    label2A.pack()   
    label3 = tk.Label(TexteP, text="El Plt maximo encontrado en la fase A vale {:.2f}".format(max(APLT)) + "\n"
                     + "El Plt maximo encontrado en la fase B vale {:.2f}".format(max(BPLT)) + "\n"
                     + "El Plt maximo encontrado en la fase C vale {:.2f}".format(max(CPLT)) + "\n" 
                     + "El Pst maximo encontrado en la fase A vale {:.2f}".format(max(APST)) + "\n"    
                     + "El Pst maximo encontrado en la fase B vale {:.2f}".format(max(BPST)) + "\n"
                     + "El Pst maximo encontrado en la fase C vale {:.2f}".format(max(CPST)), foreground = "white", background = "#2B00FA")    
    label3['font'] = f_label
    label3.pack()
 

def Flicker(fichier):
    L = extractiondonne(fichier)
    NB = []
    APLT = []
    BPLT = []
    CPLT = []
    APST = []
    BPST = []
    CPST = []
    for i in range(1,len(L)):
        NB.append(i)
    for i in range (1,len(L)):
        APLT.append(L[i][-3])
        BPLT.append(L[i][-2])
        CPLT.append(L[i][-1])
        APST.append(L[i][-6])
        BPST.append(L[i][-5])
        CPST.append(L[i][-4])
    plt.plot(NB,APLT, label = r"$Plt_A$")
    plt.legend()
    plt.plot(NB,BPLT, label = r"$Plt_B$")
    plt.legend()
    plt.plot(NB,CPLT, label = r"$Plt_C$")
    plt.legend()     
    plt.title("Variaciones de tension de largo plazo (Plt)")
    if os.path.exists(get_desktop() + "\\Nuevo Estudio" + "\\Plt.png") == False:
        plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Plt.png")
        plt.clf()
    plt.plot(NB,APST, label = r"$Pst_A$")
    plt.legend()
    plt.plot(NB,BPST, label = r"$Pst_B$")
    plt.legend()
    plt.plot(NB,CPST, label = r"$Pst_C$")
    plt.legend()     
    plt.title("Variaciones de tension de corto plazo (Pst)")
    plt.savefig(get_desktop() + "\\Nuevo Estudio" + "\\Pst.png") 
    plt.clf()       
    document.add_heading("Parte 5: Calidad de la energía ",1)
    document.add_heading("5.1 Variaciones de tension (Flicker): ",2)
    document.add_paragraph("Tenemos dos indicadores a disposicion por evaluar las variaciones de tension: "\
                           + "el indicador de variaciones de tension de largo plazo (Plt) y el indicador de variaciones"\
                           + " de tension de corto plazo (Pst).")
    document.add_paragraph("La especificacion CFE-L000045 especifica las limites siguientes: " + "\n")
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\Flicker.png", width=Inches(4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Pour le Plt
    document.add_paragraph("Este grafico mostra la evolucion del Plt de las tres fases A,B et C: " + "\n")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Plt.png", width=Inches(4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio del Plt por la fase A vale {:.2f}".format(moyenneA(APLT)))
    document.add_paragraph("- El valor promedio del Plt por la fase B vale {:.2f}".format(moyenneA(BPLT)))                       
    document.add_paragraph("- El valor promedio del Plt por la fase C vale {:.2f}".format(moyenneA(CPLT)))
    document.add_paragraph("- La mediana del Plt por la fase A vale {:.2f}".format(mediane(APLT)))
    document.add_paragraph("- La mediana del Plt por la fase B vale {:.2f}".format(mediane(BPLT)))                       
    document.add_paragraph("- La mediana del Plt por la fase C vale {:.2f}".format(mediane(CPLT)))
    document.add_paragraph("- El Plt maximo encontrado en la fase A es de {:.2f}".format(max(APLT)))
    document.add_paragraph("- El Plt maximo encontrado en la fase B es de {:.2f}".format(max(BPLT)))
    document.add_paragraph("- El Plt maximo encontrado en la fase C es de {:.2f}".format(max(CPLT)))
    document.add_paragraph("- El Plt minimo encontrado en la fase A es de {:.2f}".format(min(APLT)))
    document.add_paragraph("- El Plt minimo encontrado en la fase B es de {:.2f}".format(min(BPLT)))
    document.add_paragraph("- El Plt minimo encontrado en la fase C es de {:.2f}".format(min(CPLT)) + "\n")
    if max(APLT) > 1:
        document.add_paragraph("Cómo se puede observar, el PLT en la fase A sube la limitacion, y entonces no respeta la norma")
    if max(APLT) <= 1:
        document.add_paragraph("Cómo se puede observar, el PLT en la fase A no sube la limitacion, y entonces respeta la norma")
    if max(BPLT) > 1:
        document.add_paragraph("Cómo se puede observar, el PLT en la fase B sube la limitacion, y entonces no respeta la norma")
    if max(BPLT) <= 1:
        document.add_paragraph("Cómo se puede observar, el PLT en la fase B no sube la limitacion, y entonces respeta la norma")
    if max(CPLT) > 1:
        document.add_paragraph("Cómo se puede observar, el PLT en la fase C sube la limitacion, y entonces no respeta la norma")
    if max(CPLT) <= 1:
        document.add_paragraph("Cómo se puede observar, el PLT en la fase C no sube la limitacion, y entonces respeta la norma")   
    # Pour le Pst
    document.add_paragraph("Este grafico mostra la evolucion del Pst de las tres fases A,B et C: " + "\n")
    document.add_picture(get_desktop() + "\\Nuevo Estudio" + "\\Pst.png", width=Inches(4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Podemos observar algunas cosas en este grafico: " + "\n")
    document.add_paragraph("- El valor promedio del Pst por la fase A vale {:.2f}".format(moyenneA(APST)))
    document.add_paragraph("- El valor promedio del Pst por la fase B vale {:.2f}".format(moyenneA(BPST)))                       
    document.add_paragraph("- El valor promedio del Pst por la fase C vale {:.2f}".format(moyenneA(CPST)))
    document.add_paragraph("- La mediana del Pst por la fase A vale {:.2f}".format(mediane(APST)))
    document.add_paragraph("- La mediana del Pst por la fase B vale {:.2f}".format(mediane(BPST)))                       
    document.add_paragraph("- La mediana del Pst por la fase C vale {:.2f}".format(mediane(CPST)))
    document.add_paragraph("- El Pst maximo encontrado en la fase A es de {:.2f}".format(max(APST)))
    document.add_paragraph("- El Pst maximo encontrado en la fase B es de {:.2f}".format(max(BPST)))
    document.add_paragraph("- El Pst maximo encontrado en la fase C es de {:.2f}".format(max(CPST)))
    document.add_paragraph("- El Pst minimo encontrado en la fase A es de {:.2f}".format(min(APST)))
    document.add_paragraph("- El Pst minimo encontrado en la fase B es de {:.2f}".format(min(BPST)))
    document.add_paragraph("- El Pst minimo encontrado en la fase C es de {:.2f}".format(min(CPST)))
    if max(APST) > 0.65:
        document.add_paragraph("Cómo se puede observar, el PST en la fase A sube la limitacion, y entonces no respeta la norma")
    if max(APST) <= 0.65:
        document.add_paragraph("Cómo se puede observar, el PST en la fase A no sube la limitacion, y entonces respeta la norma")
    if max(BPST) > 0.65:
        document.add_paragraph("Cómo se puede observar, el PST en la fase B sube la limitacion, y entonces no respeta la norma")
    if max(BPST) <= 0.65:
        document.add_paragraph("Cómo se puede observar, el PST en la fase B no sube la limitacion, y entonces respeta la norma")
    if max(CPST) > 0.65:
        document.add_paragraph("Cómo se puede observar, el PST en la fase C sube la limitacion, y entonces no respeta la norma")
    if max(CPST) <= 0.65:
        document.add_paragraph("Cómo se puede observar, el PST en la fase C no sube la limitacion, y entonces respeta la norma")   
    

# En una ventana tkinter, voltea al valor promedio de Qavg, Qmin y Qmax, a sus valores medianas, al los valores
# maxima y minima encontradas y en otra ventana, desplega al grafico de Qavg, Qmin y Qmax, y a la media mobil 
# sobre 6 horas                 
    
    
def Qc(fichier):
    L = extractiondonne(fichier)
    Qfmoy = []
    Qfmax = []
    Qfmin = []
    FI = 0.98
    for i in range(1,len(L)):
        Qfmoy.append(L[i][64] - L[i][45]*tan(acos(FI)))
        Qfmax.append(L[i][65] - L[i][46]*tan(acos(FI)))
        Qfmin.append(L[i][66] - L[i][47]*tan(acos(FI)))
    plt.plot(NB,Qfmoy, label = "Qcmoy")
    plt.legend()    
    plt.plot(NB,Qfmax, label = "Qcmax")
    plt.legend()    
    plt.plot(NB,Qfmin, label = "Qcmin")
    plt.legend()

    
def mediane(L):
    A = sorted(L)
    if len(A)%2 == 1:
        return L[floor(len(L)/2)]
    else:
        return (L[floor((len(L) -1)/2)] + L[ceil((len(L) - 1)/2)])/2
    
def moyenneA(L):
    S = 0
    for i in range(len(L)):
        S+= L[i]
    return S/len(L)

#Calcule une moyenne glissante telle qu'il le font pour le covid    
def moyenneGli(L,elements): 
    Moy = [L[0]]
    for i in range(1,elements):
        Moy.append((i*Moy[i-1] + L[i])/(i + 1))
    for i in range(elements,len(L)):
        Moy.append(Moy[i-1] + (1/elements)*(L[i] - L[i - elements]))            
    return Moy

     
# Cette fonction permet d'afficher la deuxième fenêtre du programme    
def menu(fichier):
    global TN
    TN = float(saisieChamp.get())        
    global Snom
    Snom = float(saisieChamp3.get())
    global FPQ
    FPQ = float(saisieChamp4.get())
    global Icc
    Icc = float(saisieChamp2.get())    
    global Inom
    Inom = GetIM(fichier)
    myWindow.destroy()
    P(fichier)
    #plt.clf()
    S(fichier)
    #plt.clf()
    Q(fichier)
    D(fichier)
    Freq(fichier)
    TensionF(fichier)
    TensionL(fichier)
    DesT(fichier)
    DestI(fichier)
    Flicker(fichier)
    DistT(fichier)
    DistI(fichier)
    Fp(fichier)
    document.add_heading('Anexo:', 0)
    document.add_paragraph("Potencia de deformación: se puede mostrar que, considerando que la tension "\
                           + "de alimentación es sinusoidal, la potencia de deformación D se puede expresar"\
                           + " de la suiguiente manera: "    + "\n")
    document.add_picture(get_desktop() + "\\DDBAnalyzer" + "\\definition D harmo.png", width=Inches(2.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Hemos usado esta definición para calcular D, con los corrientes armonicos"\
                           + " de orden hasta 49." + "\n")
    document.add_paragraph(" Calculo de desbalances: Para calcular un desbalance, que sea de tensión o de corriente"\
                           + " el metodo es el siguiente: " + "\n")
    document.add_paragraph("- calculamos el valor promedio entre las tres fases, que llamamos Vpro")    
    document.add_paragraph("- identificamos el valor maximo entre las tres fases, que llamos Vmax")
    document.add_paragraph("- El desbalance(%) es : 100x(Vmax - Vpro)/Vpro" + "\n")
    document.save(get_desktop() + "\\Nuevo Estudio" + "\\EstudioPQBarcon.docx")    
    A = tk.Tk()
    A.title("Data analyser PQ B")
    A.configure(background="#2B00FA")
    #1er message
    f_label = font.Font(family='Times New Roman', size=20)
    f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
    label = tk.Label(text="\n"  + "¿Que quieres visualisar?" + "\n" , foreground = "white", background = "#2B00FA")
    label['font'] = f_label
    label.pack()
    # 1er bouton
    bouton1 = tk.Button(A,text='Potentia activa', command = lambda:BoutonP(fichier))
    bouton1.place(relx=0.5, rely=0.2, height=50, width=147)
    bouton1.configure(foreground="#2B00FA")
    bouton1['font'] = f_bouton
    bouton1.pack()
    # 2eme bouton
    bouton2 = tk.Button(A,text='Potentia aparente', command = lambda:BoutonS(fichier))
    bouton2.configure(foreground="#2B00FA")
    bouton2['font'] = f_bouton
    bouton2.pack()
    # un autre bouton
    bouton12 = tk.Button(A,text='Potentia reactiva', command = lambda:BoutonQ(fichier))
    bouton12.configure(foreground="#2B00FA")
    bouton12['font'] = f_bouton
    bouton12.pack()
    # 3eme bouton
    bouton3 = tk.Button(A,text='Factor de potencia', command = lambda:BoutonFp(fichier))
    bouton3.configure(foreground="#2B00FA")
    bouton3['font'] = f_bouton
    bouton3.pack()
    # 4eme bouton
    bouton4 = tk.Button(A,text='Frecuencia', command = lambda:BoutonFreq(fichier))
    bouton4.configure(foreground="#2B00FA")
    bouton4['font'] = f_bouton
    bouton4.pack()
    # 5eme bouton
    bouton5 = tk.Button(A,text='Tension de fase', command = lambda:BoutonTF(fichier))
    bouton5.configure(foreground="#2B00FA")
    bouton5['font'] = f_bouton
    bouton5.pack()
    # 6eme bouton
    bouton6 = tk.Button(A,text='Tension de linea', command = lambda:BoutonTL(fichier))
    bouton6.configure(foreground="#2B00FA")
    bouton6['font'] = f_bouton
    bouton6.pack()
    # 7eme bouton
    bouton7 = tk.Button(A,text='Desbalanceo de tension', command = lambda:BoutonDesT(fichier))
    bouton7.configure(foreground="#2B00FA")
    bouton7['font'] = f_bouton
    bouton7.pack()
    # 8eme bouton
    bouton8 = tk.Button(A,text='Desbalanceo de corriente', command = lambda:BoutonDestI(fichier))
    bouton8.configure(foreground="#2B00FA")
    bouton8['font'] = f_bouton
    bouton8.pack()
    # 9eme bouton
    bouton9 = tk.Button(A,text='Distorsion de tension', command = lambda:BoutonDistT(fichier))
    bouton9.configure(foreground="#2B00FA")
    bouton9['font'] = f_bouton
    bouton9.pack()
    # 10eme bouton
    bouton10 = tk.Button(A,text='Distorsion de corriente', command = lambda:BoutonDistI(fichier))
    bouton10.configure(foreground="#2B00FA")
    bouton10['font'] = f_bouton
    bouton10.pack()
    # 11eme Bouton
    # 10eme bouton
    bouton11 = tk.Button(A,text='Flicker', command = lambda:BoutonFlicker(fichier))
    bouton11.configure(foreground="#2B00FA")
    bouton11['font'] = f_bouton
    bouton11.pack()
    
#Ici tout ce qui est lié à Tkinter 


#C'est cette fonction qui est la coeur du programme : elle permet d'acquérir les données du fichier via
#la fonction extractiondonne, ferme la première fenètre et ouvre la seconde via la fonction menu()
def aquire():
    global fichier
    fichier = askopenfilename(filetypes=[('csv', '*.csv')])
    extractiondonne(fichier)
    interface.destroy()
    DonneIntermediaire()    
    
def DonneIntermediaire():
    global myWindow
    myWindow = tk.Tk()
    #création d'un cadre (Frame) de saisie pour le rayon
    saisieZone = tk.Frame(myWindow, borderwidth=12)
    #ajout d'un label dans le cadre
    saisieLabel = tk.Label(saisieZone, text='Tension nominal al secundario (V): ')
    #ajout dans le cadre, à gauche
    saisieLabel.pack()
    #création d'un champ de saisie dans le cadre : centré, taille de 4 caractères
    global saisieChamp
    saisieChamp = tk.Entry(saisieZone, width=10)
    #ajout dans le cadre, à droite
    saisieChamp.pack()
    #ajout du cadre dans la fenêtre
    saisieZone.pack()
    # Cadre de saisie pour Icc
    saisieZone2 = tk.Frame(myWindow, borderwidth=12)
    saisieLabel2 = tk.Label(saisieZone2, text='Corriente de corte circuito Icc (A): ')
    saisieLabel2.pack()
    global saisieChamp2
    saisieChamp2 = tk.Entry(saisieZone2, width=10)
    saisieChamp2.pack()
    saisieZone2.pack()
    # Cadre de saisie pour Snom
    saisieZone3 = tk.Frame(myWindow, borderwidth=12)
    saisieLabel3 = tk.Label(saisieZone2, text='Potencia aparente nominal del secundario (kVA): ')
    saisieLabel3.pack()
    global saisieChamp3
    saisieChamp3 = tk.Entry(saisieZone3, width=10)
    saisieChamp3.pack()
    saisieZone3.pack()  
    # Factor de potencia querido
    saisieZone4 = tk.Frame(myWindow, borderwidth=12)
    saisieLabel4 = tk.Label(saisieZone4, text='Factor de potencia querido: ')
    saisieLabel4.pack()
    global saisieChamp4
    saisieChamp4 = tk.Entry(saisieZone4, width=10)
    saisieChamp4.pack()
    saisieZone4.pack()                          
    # Boton Validar
    bouton7 = tk.Button(myWindow,text = 'Validar', command = lambda:menu(fichier))
    bouton7.configure(foreground="#2B00FA")
    bouton7['font'] = f_bouton
    bouton7.pack()
    
# Première fenêtre du programme    
interface = tk.Tk()
interface.title("Data analyser PQ B")
interface.configure(background="#2B00FA")
# 1er message dans la première fenêtre
f_label = font.Font(family='Times New Roman', size=20)
f_bouton = font.Font(family='Times New Roman', size=15, weight="bold")
label = tk.Label(text="\n"  + "¡Buenos días!" + "\n"  + "\n" + "¿Cúal estudio quieres abrir?" + "\n" , foreground = "white", background = "#2B00FA")
label['font'] = f_label
label.pack()
# 1er bouton, permet de selectionner notre étude et d'ouvrir ensuite la seconde fenêtre
bouton = tk.Button(text='Cliquea acá para abrir tu estudio', command=aquire)
bouton.place(relx=0.200, rely=0.06, height=500, width=147)
bouton.configure(foreground="#2B00FA")
bouton['font'] = f_bouton
bouton.pack(expand="yes")
    


interface.mainloop()
